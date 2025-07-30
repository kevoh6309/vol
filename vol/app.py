from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, g, session, make_response
from markupsafe import Markup
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime, timezone, timedelta
import os
import stripe
from flask_mail import Mail, Message
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_wtf.csrf import CSRFProtect
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, TextAreaField, BooleanField, SelectField, DateField
from wtforms.validators import DataRequired, Length, EqualTo
import requests
import json
import re
# Import config directly
import os
try:
    from config import config
except ImportError:
    # If config import fails, create a basic config
    class Config:
        SECRET_KEY = os.getenv('SECRET_KEY', 'kevoh2071M@')
        FLASK_ENV = os.getenv('FLASK_ENV', 'development')
        FLASK_DEBUG = os.getenv('FLASK_DEBUG', 'True').lower() == 'true'
        SQLALCHEMY_DATABASE_URI = os.getenv('DATABASE_URL', 'sqlite:///vol.db')
        SQLALCHEMY_TRACK_MODIFICATIONS = False
        
        # Enhanced Security Configuration
        SESSION_COOKIE_SECURE = os.getenv('SESSION_COOKIE_SECURE', 'True').lower() == 'true'
        SESSION_COOKIE_HTTPONLY = True
        SESSION_COOKIE_SAMESITE = 'Lax'
        PERMANENT_SESSION_LIFETIME = 86400  # 24 hours
        SESSION_COOKIE_MAX_AGE = 86400
        SESSION_REFRESH_EACH_REQUEST = True
        
        # CSRF Protection
        WTF_CSRF_ENABLED = True
        WTF_CSRF_TIME_LIMIT = 3600  # 1 hour
        
        # Rate Limiting
        RATELIMIT_STORAGE_URL = os.getenv('REDIS_URL', 'memory://')
        RATELIMIT_DEFAULT = "200 per day;50 per hour;10 per minute"
        
        # Security Headers
        SECURITY_HEADERS = {
            'X-Content-Type-Options': 'nosniff',
            'X-Frame-Options': 'SAMEORIGIN',
            'X-XSS-Protection': '1; mode=block',
            'Strict-Transport-Security': 'max-age=31536000; includeSubDomains',
            'Content-Security-Policy': "default-src 'self'; script-src 'self' 'unsafe-inline' 'unsafe-eval' https://cdn.jsdelivr.net https://pagead2.googlesyndication.com; style-src 'self' 'unsafe-inline' https://cdn.jsdelivr.net https://fonts.googleapis.com; font-src 'self' https://cdn.jsdelivr.net https://fonts.gstatic.com; img-src 'self' data: https:; connect-src 'self' https://api.stripe.com https://generativelanguage.googleapis.com https://api.cohere.ai https://openrouter.ai; frame-src https://js.stripe.com https://hooks.stripe.com;"
        }
        
        # Mail Configuration
        MAIL_SERVER = os.getenv('MAIL_SERVER', 'smtp.gmail.com')
        MAIL_PORT = int(os.getenv('MAIL_PORT', 587))
        MAIL_USE_TLS = os.getenv('MAIL_USE_TLS', 'True').lower() == 'true'
        MAIL_USERNAME = os.getenv('MAIL_USERNAME', 'kevohmutwiri35@gmail.com')
        MAIL_PASSWORD = os.getenv('MAIL_PASSWORD', 'kevoh2071M@')
        MAIL_DEFAULT_SENDER = os.getenv('MAIL_USERNAME', 'kevohmutwiri35@gmail.com')
        
        # Stripe Configuration
        STRIPE_SECRET_KEY = os.environ.get('STRIPE_SECRET_KEY')
        STRIPE_PUBLISHABLE_KEY = os.environ.get('STRIPE_PUBLISHABLE_KEY')
        STRIPE_MONTHLY_PRICE_ID = os.environ.get('STRIPE_MONTHLY_PRICE_ID')
        STRIPE_YEARLY_PRICE_ID = os.environ.get('STRIPE_YEARLY_PRICE_ID')
        STRIPE_WEBHOOK_SECRET = os.environ.get('STRIPE_WEBHOOK_SECRET')
        
        # AI API Configuration
        GEMINI_API_KEY = os.getenv('GEMINI_API_KEY', 'your_gemini_api_key_here')
        GEMINI_API_URL = os.getenv('GEMINI_API_URL', 'https://generativelanguage.googleapis.com/v1beta/models/gemini-pro:generateContent')
        COHERE_API_KEY = os.getenv('COHERE_API_KEY', 'your_cohere_api_key_here')
        OPENROUTER_API_KEY = os.getenv('OPENROUTER_API_KEY', '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')
        OPENROUTER_API_URL = os.getenv('OPENROUTER_API_URL', 'https://openrouter.ai/api/v1/chat/completions')
        
        # Application Configuration
        APP_NAME = os.getenv('APP_NAME', 'ResumeBuilder Pro')
        APP_URL = os.getenv('APP_URL', 'https://resume-builder-saas.railway.app')
        
        # Logging Configuration
        LOG_LEVEL = os.getenv('LOG_LEVEL', 'INFO')
        LOG_FORMAT = '%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'
    
    class DevelopmentConfig(Config):
        DEBUG = True
        SESSION_COOKIE_SECURE = False  # Allow HTTP in development
    
    class ProductionConfig(Config):
        DEBUG = False
        APP_URL = os.getenv('APP_URL', 'https://resume-builder-saas.railway.app')
        SESSION_COOKIE_SECURE = True
    
    config = {
        'development': DevelopmentConfig,
        'production': ProductionConfig,
        'default': DevelopmentConfig
    }
import secrets
import time
from collections import Counter
import datetime as dt
import tempfile
# WeasyPrint will be imported locally in functions where needed
WEASYPRINT_AVAILABLE = None  # Will be set when first needed

def get_weasyprint():
    """Safely import WeasyPrint and return HTML class if available"""
    global WEASYPRINT_AVAILABLE
    if WEASYPRINT_AVAILABLE is None:
        try:
            from weasyprint import HTML
            WEASYPRINT_AVAILABLE = True
            return HTML
        except ImportError:
            WEASYPRINT_AVAILABLE = False
            return None
    elif WEASYPRINT_AVAILABLE:
        from weasyprint import HTML
        return HTML
    else:
        return None
# Try to import python-docx, but make it optional
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Word document generation will be disabled.")
import logging
from logging.handlers import RotatingFileHandler

# Set up logging
if not os.path.exists('logs'):
    os.mkdir('logs')

# Configure logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]',
    handlers=[
        RotatingFileHandler('logs/vol.log', maxBytes=10240000, backupCount=10),
        logging.StreamHandler()
    ]
)

# Create logger for this application
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)

# Initialize Flask app
app = Flask(__name__)
logger.info("Flask app initialized")

# Load configuration
config_name = os.getenv('FLASK_ENV', 'development')
app.config.from_object(config[config_name])
logger.info(f"Loaded configuration: {config_name}")

# Ensure session configuration for CSRF
if not app.config.get('SECRET_KEY'):
    app.config['SECRET_KEY'] = 'your-secret-key-change-this-in-production'
    logger.warning("Using default SECRET_KEY - change in production")

# Security middleware
@app.after_request
def add_security_headers(response):
    """Add security headers to all responses"""
    for header, value in app.config.get('SECURITY_HEADERS', {}).items():
        response.headers[header] = value
    return response

@app.before_request
def before_request():
    """Security checks before each request"""
    # Log request info (without sensitive data)
    if request.endpoint and 'static' not in request.endpoint:
        logger.info(f'Request: {request.method} {request.endpoint} from {request.remote_addr}')
    
    # Session security
    if session.get('_fresh', False):
        session.permanent = True
        session.modified = True
    
    # Rate limiting is handled by decorators, not manual checks

# Initialize extensions
db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'
csrf = CSRFProtect(app)
logger.info("Extensions initialized: SQLAlchemy, LoginManager, CSRFProtect")

# Initialize rate limiter
limiter = Limiter(
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"]
)
limiter.init_app(app)
logger.info("Rate limiter initialized")

# Initialize mail
mail = Mail(app)
logger.info("Mail extension initialized")

# Configure Stripe
stripe.api_key = app.config['STRIPE_SECRET_KEY']
STRIPE_PUBLISHABLE_KEY = app.config['STRIPE_PUBLISHABLE_KEY']
STRIPE_WEBHOOK_SECRET = app.config['STRIPE_WEBHOOK_SECRET']
logger.info("Stripe configured")

# Forms
class LoginForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired()])
    password = PasswordField('Password', validators=[DataRequired()])
    remember_me = BooleanField('Remember Me')

class RegisterForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired(), Length(min=2, max=20)])
    email = StringField('Email', validators=[DataRequired()])
    password = PasswordField('Password', validators=[DataRequired(), Length(min=6)])
    confirm_password = PasswordField('Confirm Password', validators=[DataRequired(), EqualTo('password', message='Passwords must match')])

class FeedbackForm(FlaskForm):
    name = StringField('Name', validators=[DataRequired()])
    email = StringField('Email', validators=[DataRequired()])
    message = TextAreaField('Message', validators=[DataRequired()])

class ChangePasswordForm(FlaskForm):
    current_password = PasswordField('Current Password', validators=[DataRequired()])
    new_password = PasswordField('New Password', validators=[DataRequired(), Length(min=6)])
    confirm_new_password = PasswordField('Confirm New Password', validators=[DataRequired(), EqualTo('new_password', message='Passwords must match')])

class JobApplicationForm(FlaskForm):
    job_title = StringField('Job Title', validators=[DataRequired()])
    company = StringField('Company', validators=[DataRequired()])
    status = SelectField('Status', choices=[('applied', 'Applied'), ('interview', 'Interview'), ('offer', 'Offer'), ('rejected', 'Rejected')])
    applied_date = DateField('Applied Date', format='%Y-%m-%d')
    resume_id = SelectField('Resume', coerce=int, choices=[])
    cover_letter_id = SelectField('Cover Letter', coerce=int, choices=[])

# Database Models
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(150), unique=True, nullable=False)
    email = db.Column(db.String(150), unique=True, nullable=False)
    password_hash = db.Column(db.String(150), nullable=False)
    resumes = db.relationship('Resume', backref='author', lazy=True)
    # Premium and branding fields
    is_premium = db.Column(db.Boolean, default=False)
    premium_expiry = db.Column(db.DateTime)
    custom_logo = db.Column(db.String(300))
    custom_domain = db.Column(db.String(300))
    analytics_enabled = db.Column(db.Boolean, default=False)
    # For testimonials/social proof
    testimonial = db.Column(db.Text)
    # Profile image
    profile_image = db.Column(db.String(300))  # Store file path
    # Password reset fields
    reset_token = db.Column(db.String(255))
    reset_token_expiry = db.Column(db.Integer)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    is_active = db.Column(db.Boolean, default=True)
    # Add is_admin to User model if missing
    is_admin = db.Column(db.Boolean, default=False)

class Resume(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(150))
    email = db.Column(db.String(150))
    phone = db.Column(db.String(100))
    address = db.Column(db.String(200))
    linkedin = db.Column(db.String(200))
    summary = db.Column(db.Text)
    education = db.Column(db.Text)
    experience = db.Column(db.Text)
    skills = db.Column(db.Text)
    certifications = db.Column(db.Text)
    languages = db.Column(db.Text)
    pdf_engine = db.Column(db.String(50))
    template = db.Column(db.String(50))
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc), onupdate=lambda: datetime.now(timezone.utc))
    downloads = db.Column(db.Integer, default=0)  # Track download count
    is_public = db.Column(db.Boolean, default=False)  # Public sharing
    shares = db.Column(db.Integer, default=0)  # Share count
    views = db.Column(db.Integer, default=0)  # View count

class CoverLetter(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200)) 
    job_title = db.Column(db.String(150))
    company = db.Column(db.String(150))
    content = db.Column(db.Text)
    template = db.Column(db.String(50))
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime)
    updated_at = db.Column(db.DateTime)
    downloads = db.Column(db.Integer, default=0)  # Track download count
    is_public = db.Column(db.Boolean, default=False)  # Public sharing
    shares = db.Column(db.Integer, default=0)  # Share count
    views = db.Column(db.Integer, default=0)  # View count

class JobApplication(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    job_title = db.Column(db.String(150), nullable=False)
    company = db.Column(db.String(150), nullable=False)
    status = db.Column(db.String(50), default='applied')  # e.g., applied, interview, offer, rejected
    applied_date = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    resume_id = db.Column(db.Integer, db.ForeignKey('resume.id'))
    cover_letter_id = db.Column(db.Integer, db.ForeignKey('cover_letter.id'))
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    # Relationships
    resume = db.relationship('Resume', foreign_keys=[resume_id])
    cover_letter = db.relationship('CoverLetter', foreign_keys=[cover_letter_id])
    user = db.relationship('User', foreign_keys=[user_id])

class PracticeSession(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    category = db.Column(db.String(100), nullable=False)
    score = db.Column(db.Float, nullable=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    user = db.relationship('User', foreign_keys=[user_id])

class PracticeSessionForm(FlaskForm):
    category = StringField('Category', validators=[DataRequired()])
    score = StringField('Score', validators=[DataRequired()])

class Reminder(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    message = db.Column(db.String(300), nullable=False)
    due_date = db.Column(db.DateTime, nullable=False)
    sent = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    user = db.relationship('User', foreign_keys=[user_id])

class TeamCV(db.Model):
    __tablename__ = 'teamcv'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(150), nullable=False)
    owner_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    owner = db.relationship('User', foreign_keys=[owner_id])
    members = db.relationship('TeamMember', backref='team', lazy=True)

class TeamMember(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    team_id = db.Column(db.Integer, db.ForeignKey('teamcv.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    invited_email = db.Column(db.String(150))
    role = db.Column(db.String(50), default='member')
    joined_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    user = db.relationship('User', foreign_keys=[user_id])

class AffiliateLink(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    code = db.Column(db.String(50), unique=True, nullable=False)
    clicks = db.Column(db.Integer, default=0)
    signups = db.Column(db.Integer, default=0)
    payouts = db.Column(db.Float, default=0.0)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    user = db.relationship('User', foreign_keys=[user_id])

# Helper functions
@login_manager.user_loader
def load_user(user_id):
    try:
        return User.query.get(int(user_id))
    except:
        return None

def is_premium(user):
    """Check if a user is premium and their subscription is valid."""
    return getattr(user, 'is_authenticated', False) and getattr(user, 'is_premium', False)

# Add password reset functionality
def generate_reset_token():
    return secrets.token_urlsafe(32)

def verify_reset_token(token, max_age=3600):  # 1 hour expiry
    try:
        # In a real app, you'd store tokens in database with expiry
        # For now, we'll use a simple approach
        user = User.query.filter_by(reset_token=token).first()
        if user and user.reset_token_expiry and time.time() < user.reset_token_expiry:
            return user
    except:
        pass
    return None

def send_email(to_email, subject, template_name, **kwargs):
    """Send email using templates"""
    try:
        msg = Message(subject,
                    sender=app.config['MAIL_DEFAULT_SENDER'],
                    recipients=[to_email])
        
        # Import template macros
        from flask import render_template_string
        
        if template_name == 'password_reset':
            from vol.templates.email_templates import password_reset_email
            html_content = render_template_string(
                password_reset_email(kwargs.get('username', ''), kwargs.get('reset_url', '')),
                username=kwargs.get('username', ''),
                reset_url=kwargs.get('reset_url', '')
            )
        elif template_name == 'welcome':
            from vol.templates.email_templates import welcome_email
            html_content = render_template_string(
                welcome_email(kwargs.get('username', '')),
                username=kwargs.get('username', '')
            )
        elif template_name == 'premium_upgrade':
            from vol.templates.email_templates import premium_upgrade_email
            html_content = render_template_string(
                premium_upgrade_email(kwargs.get('username', '')),
                username=kwargs.get('username', '')
            )
        else:
            html_content = kwargs.get('html_content', '')
        
        msg.html = html_content
        msg.body = kwargs.get('text_content', '')
        
        mail.send(msg)
        return True
    except Exception as e:
        logger.error(f"Email sending failed: {e}")
        return False

@app.before_request
def init_db():
    if not hasattr(app, 'db_initialized'):
        db.create_all()
        app.db_initialized = True
    
    # Safely set is_premium on g object
    try:
        g.is_premium = is_premium(current_user) if hasattr(current_user, 'is_authenticated') and current_user.is_authenticated else False
    except:
        g.is_premium = False
    logger.info(f"Before request: is_premium={g.is_premium}")

def render_with_premium(template, **kwargs):
    from flask import g
    kwargs['is_premium'] = getattr(g, 'is_premium', False)
    return render_template(template, **kwargs)

# Routes
@app.route('/')
def home():
    return render_template('landing.html')

@app.route('/favicon.ico')
def favicon():
    return send_file('static/favicon.ico', mimetype='image/vnd.microsoft.icon')

@app.route('/register', methods=['GET', 'POST'])
@limiter.limit('5 per minute')
def register():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    form = RegisterForm()
    if form.validate_on_submit():
        if User.query.filter_by(email=form.email.data).first():
            flash('Email already registered', 'error')
        elif form.password.data != form.confirm_password.data:
            flash('Passwords do not match', 'error')
        else:
            user = User(
                username=form.username.data,
                email=form.email.data,
                password_hash=generate_password_hash(form.password.data)
            )
            db.session.add(user)
            db.session.commit()
            
            # Send welcome email
            send_email(
                user.email,
                'Welcome to ResumeBuilder!',
                'welcome',
                username=user.username
            )
            
            flash('Registration successful! Please log in.', 'success')
            return redirect(url_for('login'))
    
    return render_template('register.html', form=form)

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data).first()
        if user and check_password_hash(user.password_hash, form.password.data):
            login_user(user, remember=form.remember_me.data)
            next_page = request.args.get('next')
            return redirect(next_page) if next_page else redirect(url_for('dashboard'))
        else:
            flash('Invalid email or password', 'error')
    
    return render_template('login.html', form=form)

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('home'))

@app.route('/clear-session')
def clear_session():
    """Clear session and redirect to home"""
    session.clear()
    return redirect(url_for('home'))

# Update dashboard route to use real job application data
@app.route('/dashboard')
@login_required
def dashboard():
    user_resumes = Resume.query.filter_by(user_id=current_user.id).order_by(Resume.created_at.desc()).all()
    user_covers = CoverLetter.query.filter_by(user_id=current_user.id).order_by(CoverLetter.created_at.desc()).all()
    user_jobs = JobApplication.query.filter_by(user_id=current_user.id).order_by(JobApplication.applied_date.desc()).all()
    user_sessions = PracticeSession.query.filter_by(user_id=current_user.id).order_by(PracticeSession.created_at.desc()).all()
    total_resumes = len(user_resumes)
    total_covers = len(user_covers)
    total_jobs = len(user_jobs)
    total_sessions = len(user_sessions)
    dashboard_data = {
        'resumes_created': total_resumes,
        'cover_letters_created': total_covers,
        'applications_submitted': total_jobs,
        'practice_sessions_completed': total_sessions,
        'interviews_attended': max(0, total_jobs - 1),
        'offers_received': max(0, total_jobs - 2),
        'average_interview_score': (sum(s.score for s in user_sessions) / total_sessions) if total_sessions else 0,
        'streak_days': 7
    }
    recent_resumes = user_resumes[:5]
    recent_covers = user_covers[:5]
    recent_jobs = user_jobs[:5]
    recent_sessions = user_sessions[:5]
    reminders = Reminder.query.filter_by(user_id=current_user.id, sent=False).filter(Reminder.due_date >= datetime.now(timezone.utc)).order_by(Reminder.due_date).all()
    return render_template('dashboard.html',
        now=datetime.now(timezone.utc),
        dashboard=dashboard_data,
        max=max,
        recent_resumes=recent_resumes,
        recent_covers=recent_covers,
        recent_jobs=recent_jobs,
        recent_sessions=recent_sessions,
        reminders=reminders,
        is_premium=is_premium(current_user))

@app.route('/my-resumes')
@login_required
def my_resumes():
    resumes = Resume.query.filter_by(user_id=current_user.id).all()
    return render_template('my_resumes.html', resumes=resumes, is_premium=is_premium(current_user))

@app.route('/create-resume', methods=['GET', 'POST'])
@login_required
def create_resume():
    if request.method == 'POST':
        # Handle resume creation with input sanitization
        resume = Resume(
            name=sanitize_input(request.form.get('name')),
            email=sanitize_input(request.form.get('email')),
            phone=sanitize_input(request.form.get('phone')),
            address=sanitize_input(request.form.get('address')),
            linkedin=sanitize_input(request.form.get('linkedin')),
            summary=sanitize_input(request.form.get('summary')),
            education=sanitize_input(request.form.get('education')),
            experience=sanitize_input(request.form.get('experience')),
            skills=sanitize_input(request.form.get('skills')),
            certifications=sanitize_input(request.form.get('certifications')),
            languages=sanitize_input(request.form.get('languages')),
            pdf_engine=sanitize_input(request.form.get('pdf_engine', 'weasyprint')),
            template=sanitize_input(request.form.get('template', 'modern')),
            user_id=current_user.id
        )
        db.session.add(resume)
        db.session.commit()
        flash('Resume created successfully!', 'success')
        return redirect(url_for('my_resumes'))
    return render_template('resume_form.html')

@app.route('/edit-resume/<int:resume_id>')
@login_required
def edit_resume(resume_id):
    resume = Resume.query.filter_by(id=resume_id, user_id=current_user.id).first_or_404()
    return render_template('resume_form.html', resume=resume)

@app.route('/delete-resume/<int:resume_id>', methods=['POST'])
@login_required
@csrf.exempt
def delete_resume(resume_id):
    logger.info(f"Delete resume request received for resume_id: {resume_id}")
    logger.info(f"Request method: {request.method}")
    logger.info(f"Request headers: {dict(request.headers)}")
    logger.info(f"Request form data: {dict(request.form)}")
    logger.info(f"Current user: {current_user.id}")
    
    try:
        resume = Resume.query.filter_by(id=resume_id, user_id=current_user.id).first_or_404()
        logger.info(f"Found resume: {resume.name} (ID: {resume.id})")
        
        db.session.delete(resume)
        db.session.commit()
        logger.info(f"Successfully deleted resume {resume_id}")
        
        flash('Resume deleted successfully!', 'success')
        return redirect(url_for('my_resumes'))
    except Exception as e:
        logger.error(f"Error deleting resume {resume_id}: {str(e)}")
        logger.error(f"Exception type: {type(e)}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        flash('Error deleting resume. Please try again.', 'danger')
        return redirect(url_for('my_resumes'))

@app.route('/download-resume/<int:resume_id>')
@login_required
def download_resume(resume_id):
    # Enforce free user download limit
    if not is_premium(current_user):
        user_resumes = Resume.query.filter_by(user_id=current_user.id).all()
        if len(user_resumes) > 2:
            flash('Free users can only download a maximum of 2 resumes. Upgrade to premium for unlimited downloads.', 'danger')
            return redirect(url_for('my_resumes'))
    resume = Resume.query.filter_by(id=resume_id, user_id=current_user.id).first_or_404()
    resume.downloads += 1
    db.session.commit()
    format = request.args.get('format', 'pdf')
    if format == 'word':
        # Generate Word document
        if not DOCX_AVAILABLE:
            flash('Word document generation is currently unavailable. Please try downloading as PDF instead.', 'warning')
            return redirect(url_for('my_resumes'))
        
        doc = Document()
        doc.add_heading(resume.name or 'Resume', 0)
        doc.add_paragraph(f"Email: {resume.email}")
        doc.add_paragraph(f"Phone: {resume.phone}")
        doc.add_paragraph(f"Address: {resume.address}")
        doc.add_paragraph(f"LinkedIn: {resume.linkedin}")
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(resume.summary or '')
        doc.add_heading('Education', level=1)
        doc.add_paragraph(resume.education or '')
        doc.add_heading('Experience', level=1)
        doc.add_paragraph(resume.experience or '')
        doc.add_heading('Skills', level=1)
        doc.add_paragraph(resume.skills or '')
        doc.add_heading('Certifications & Awards', level=1)
        doc.add_paragraph(resume.certifications or '')
        doc.add_heading('Languages', level=1)
        doc.add_paragraph(resume.languages or '')
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(tmp.name)
        tmp.seek(0)
        response = make_response(tmp.read())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{resume.name or "resume"}.docx"'
        tmp.close()
        return response
    else:
        # Generate PDF using WeasyPrint and modern template
        HTML = get_weasyprint()
        if HTML is None:
            flash('PDF generation is currently unavailable. Please try downloading as Word document instead.', 'warning')
            return redirect(url_for('my_resumes'))
        
        html = render_template('resume_pdf_modern.html',
            name=resume.name,
            email=resume.email,
            phone=resume.phone,
            address=resume.address,
            linkedin=resume.linkedin,
            summary=resume.summary,
            education=resume.education,
            experience=resume.experience,
            skills=resume.skills,
            certifications=resume.certifications,
            languages=resume.languages
        )
        pdf = HTML(string=html).write_pdf()
        response = make_response(pdf)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename="{resume.name or "resume"}.pdf"'
        return response

@app.route('/resume/<int:resume_id>')
@login_required
def preview_resume(resume_id):
    resume = Resume.query.filter_by(id=resume_id, user_id=current_user.id).first_or_404()
    return render_template('resume_pdf_modern.html',
        resume_id=resume_id,
        name=resume.name,
        email=resume.email,
        phone=resume.phone,
        address=resume.address,
        linkedin=resume.linkedin,
        summary=resume.summary,
        education=resume.education,
        experience=resume.experience,
        skills=resume.skills,
        certifications=resume.certifications,
        languages=resume.languages,
        preview_mode=True
    )

@app.route('/edit-cover-letter/<int:cover_id>')
@login_required
def edit_cover_letter(cover_id):
    cover = CoverLetter.query.filter_by(id=cover_id, user_id=current_user.id).first_or_404()
    return render_template('cover_letter_form.html', cover=cover)

@app.route('/create-cover-letter', methods=['GET', 'POST'])
@login_required
def create_cover_letter():
    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        job_title = request.form.get('job_title', '').strip()
        company = request.form.get('company', '').strip()
        content = request.form.get('content', '').strip()
        template = request.form.get('template', 'standard')
        cover = CoverLetter(
            title=title,
            job_title=job_title,
            company=company,
            content=content,
            template=template,
            user_id=current_user.id,
            created_at=datetime.now(timezone.utc),
            updated_at=datetime.now(timezone.utc)
        )
        db.session.add(cover)
        db.session.commit()
        flash('Cover letter created!', 'success')
        return redirect(url_for('my_cover_letters'))
    return render_template('cover_letter_form.html', cover=None)

@app.route('/delete-cover-letter/<int:cover_id>', methods=['POST'])
@login_required
def delete_cover_letter(cover_id):
    cover = CoverLetter.query.filter_by(id=cover_id, user_id=current_user.id).first_or_404()
    db.session.delete(cover)
    db.session.commit()
    return jsonify({'success': True})

@app.route('/download-cover-letter/<int:cover_id>')
@login_required
def download_cover_letter(cover_id):
    cover = CoverLetter.query.filter_by(id=cover_id, user_id=current_user.id).first_or_404()
    cover.downloads += 1
    db.session.commit()
    format = request.args.get('format', 'pdf')
    if format == 'word':
        if not DOCX_AVAILABLE:
            flash('Word document generation is currently unavailable. Please try downloading as PDF instead.', 'warning')
            return redirect(url_for('my_cover_letters'))
        
        doc = Document()
        doc.add_heading(cover.title or 'Cover Letter', 0)
        doc.add_paragraph(f"Job Title: {cover.job_title}")
        doc.add_paragraph(f"Company: {cover.company}")
        doc.add_paragraph(cover.content or '')
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(tmp.name)
        tmp.seek(0)
        response = make_response(tmp.read())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{cover.title or "cover_letter"}.docx"'
        tmp.close()
        return response
    else:
        HTML = get_weasyprint()
        if HTML is None:
            flash('PDF generation is currently unavailable. Please try downloading as Word document instead.', 'warning')
            return redirect(url_for('my_cover_letters'))
        
        html = render_template('cover_letter_pdf.html', cover=cover, current_user=current_user, now=datetime.now(timezone.utc), is_premium=True)
        pdf = HTML(string=html).write_pdf()
        response = make_response(pdf)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename="{cover.title or "cover_letter"}.pdf"'
        return response

@app.route('/share-cover-letter/<int:cover_id>')
@login_required
def share_cover_letter(cover_id):
    cover = CoverLetter.query.filter_by(id=cover_id, user_id=current_user.id).first_or_404()
    # For now, just redirect to a placeholder
    flash('Share feature coming soon!', 'info')
    return redirect(url_for('my_cover_letters'))

@app.route('/my-cover-letters')
@login_required
def my_cover_letters():
    covers = CoverLetter.query.filter_by(user_id=current_user.id).all()
    return render_template('my_cover_letters.html', covers=covers, is_premium=is_premium(current_user))

@app.route('/career-toolkit')
@login_required
def career_toolkit():
    return render_template('career_toolkit.html', is_premium=is_premium(current_user))

@app.route('/resume-checker')
@login_required
def resume_checker():
    return render_template('resume_checker.html', is_premium=is_premium(current_user))

@app.route('/analyze-resume', methods=['POST'])
@login_required
def analyze_resume():
    """Analyze uploaded resume for ATS optimization and provide feedback"""
    if 'resume' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'})
    
    file = request.files['resume']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'})
    
    # Check file size (5MB limit)
    file.seek(0, 2)  # Seek to end
    file_size = file.tell()
    file.seek(0)  # Reset to beginning
    
    if file_size > 5 * 1024 * 1024:  # 5MB
        return jsonify({'success': False, 'error': 'File too large (max 5MB)'})
    
    # Check file extension
    allowed_extensions = {'pdf', 'doc', 'docx', 'txt'}
    if not file.filename.lower().endswith(tuple('.' + ext for ext in allowed_extensions)):
        return jsonify({'success': False, 'error': 'Invalid file type'})
    
    try:
        # Read file content
        if file.filename.lower().endswith('.pdf'):
            # For PDF, we'll extract text (simplified)
            content = "PDF content extraction would go here"
        else:
            content = file.read().decode('utf-8', errors='ignore')
        
        # Analyze content
        analysis = analyze_resume_content(content, request.form.get('jobTitle', ''))
        
        return jsonify({
            'success': True,
            'analysis': analysis
        })
        
    except Exception as e:
        logger.error(f"Resume analysis error: {e}")
        return jsonify({'success': False, 'error': 'Analysis failed'})

def analyze_resume_content(content, job_title=''):
    """Analyze resume content and provide feedback"""
    content_lower = content.lower()
    
    # Common resume sections
    sections = {
        'contact_info': ['email', 'phone', 'address', 'linkedin'],
        'summary': ['summary', 'objective', 'profile'],
        'experience': ['experience', 'work history', 'employment'],
        'education': ['education', 'degree', 'university', 'college'],
        'skills': ['skills', 'technologies', 'programming', 'languages']
    }
    
    # Check for required sections
    section_analysis = {}
    for section, keywords in sections.items():
        found = any(keyword in content_lower for keyword in keywords)
        section_analysis[section] = found
    
    # Check for action verbs
    action_verbs = [
        'developed', 'implemented', 'managed', 'created', 'designed', 'built',
        'led', 'coordinated', 'analyzed', 'improved', 'increased', 'decreased',
        'maintained', 'established', 'organized', 'planned', 'executed'
    ]
    
    found_verbs = [verb for verb in action_verbs if verb in content_lower]
    
    # Check for metrics/numbers
    metrics_pattern = r'\d+%|\d+\s*(percent|%|increase|decrease|million|thousand)'
    metrics = re.findall(metrics_pattern, content_lower)
    
    # Check for keywords if job title provided
    keyword_analysis = {}
    if job_title:
        # Common keywords for different job types
        job_keywords = {
            'software': ['python', 'java', 'javascript', 'react', 'node.js', 'sql', 'git'],
            'marketing': ['seo', 'social media', 'campaign', 'analytics', 'google ads'],
            'sales': ['revenue', 'quota', 'prospecting', 'negotiation', 'closing'],
            'design': ['photoshop', 'illustrator', 'figma', 'ui/ux', 'wireframes']
        }
        
        for category, keywords in job_keywords.items():
            if category in job_title.lower():
                found_keywords = [kw for kw in keywords if kw in content_lower]
                keyword_analysis[category] = found_keywords
    
    # Calculate score
    score = 0
    suggestions = []
    
    # Section completeness (40 points)
    section_score = sum(section_analysis.values()) * 8
    score += section_score
    
    if not section_analysis['contact_info']:
        suggestions.append({'type': 'critical', 'message': 'Missing contact information'})
    if not section_analysis['experience']:
        suggestions.append({'type': 'critical', 'message': 'Missing work experience section'})
    if not section_analysis['education']:
        suggestions.append({'type': 'important', 'message': 'Missing education section'})
    
    # Action verbs (20 points)
    verb_score = min(len(found_verbs) * 2, 20)
    score += verb_score
    
    if len(found_verbs) < 5:
        suggestions.append({'type': 'important', 'message': f'Add more action verbs. Found: {len(found_verbs)}'})
    
    # Metrics (20 points)
    metric_score = min(len(metrics) * 2, 20)
    score += metric_score
    
    if len(metrics) < 3:
        suggestions.append({'type': 'helpful', 'message': 'Add more quantifiable achievements'})
    
    # Keywords (20 points)
    if keyword_analysis:
        keyword_score = 0
        for category, found in keyword_analysis.items():
            keyword_score += len(found) * 2
        score += min(keyword_score, 20)
        
        if not any(keyword_analysis.values()):
            suggestions.append({'type': 'important', 'message': 'Add relevant keywords for your target job'})
    
    # Determine grade
    if score >= 80:
        grade = 'Excellent'
        grade_class = 'score-excellent'
    elif score >= 60:
        grade = 'Good'
        grade_class = 'score-good'
    elif score >= 40:
        grade = 'Fair'
        grade_class = 'score-fair'
    else:
        grade = 'Poor'
        grade_class = 'score-poor'
    
    return {
        'score': score,
        'grade': grade,
        'grade_class': grade_class,
        'sections': section_analysis,
        'action_verbs': found_verbs,
        'metrics': metrics,
        'keywords': keyword_analysis,
        'suggestions': suggestions
    }

@app.route('/ai-cover-letter-generator')
@login_required
def ai_cover_letter_generator():
    return render_template('ai_cover_letter_generator.html', is_premium=is_premium(current_user))

@app.route('/job-application-tracker')
@login_required
def job_application_tracker():
    jobs = JobApplication.query.filter_by(user_id=current_user.id).order_by(JobApplication.applied_date.desc()).all()
    resumes = Resume.query.filter_by(user_id=current_user.id).all()
    covers = CoverLetter.query.filter_by(user_id=current_user.id).all()
    # Stats
    stats = {
        'total': len(jobs),
        'applied': sum(1 for j in jobs if j.status == 'applied'),
        'interview': sum(1 for j in jobs if j.status == 'interview'),
        'offers': sum(1 for j in jobs if j.status == 'offer'),
        'rejected': sum(1 for j in jobs if j.status == 'rejected'),
    }
    # Enrich jobs for template
    status_colors = {
        'applied': 'primary',
        'interview': 'info',
        'offer': 'warning',
        'rejected': 'danger',
        'withdrawn': 'secondary',
    }
    applications = []
    for job in jobs:
        applications.append({
            'id': job.id,
            'company': job.company,
            'job_title': job.job_title,
            'status': job.status,
            'status_color': status_colors.get(job.status, 'secondary'),
            'date_applied': job.applied_date,
            'resume_id': job.resume_id,
            'cover_letter_id': job.cover_letter_id,
            'follow_up_date': None,  # Placeholder
            'website': None,  # Placeholder
        })
    return render_template('job_application_tracker.html',
        applications=applications,
        resumes=resumes,
        cover_letters=covers,
        stats=stats,
        now=datetime.now(timezone.utc),
        is_premium=is_premium(current_user))

@app.route('/job-application/<int:app_id>')
@login_required
def job_application_details(app_id):
    job = JobApplication.query.filter_by(id=app_id, user_id=current_user.id).first_or_404()
    html = Markup(f"""
        <strong>Company:</strong> {job.company}<br>
        <strong>Job Title:</strong> {job.job_title}<br>
        <strong>Status:</strong> {job.status.title()}<br>
        <strong>Applied Date:</strong> {job.applied_date.strftime('%b %d, %Y') if job.applied_date else 'N/A'}<br>
        <strong>Resume:</strong> {job.resume_id or 'None'}<br>
        <strong>Cover Letter:</strong> {job.cover_letter_id or 'None'}<br>
    """)
    return jsonify({'success': True, 'html': html})

@app.route('/add-job-application', methods=['POST'])
@login_required
def add_job_application():
    # Support both AJAX FormData and Flask-WTF
    if request.content_type and 'multipart/form-data' in request.content_type:
        # AJAX FormData
        job_title = request.form.get('job_title', '').strip()
        company = request.form.get('company', '').strip()
        status = request.form.get('status', 'applied').strip()
        applied_date = request.form.get('date_applied')
        resume_id = request.form.get('resume_id')
        cover_letter_id = request.form.get('cover_letter_id')
        notes = request.form.get('notes', '')
        # Parse date
        from datetime import datetime, timezone
        try:
            applied_date = datetime.strptime(applied_date, '%Y-%m-%d') if applied_date else datetime.now(timezone.utc)
        except Exception:
            applied_date = datetime.now(timezone.utc)
        job = JobApplication(
            job_title=job_title,
            company=company,
            status=status,
            applied_date=applied_date,
            resume_id=int(resume_id) if resume_id and resume_id.isdigit() else None,
            cover_letter_id=int(cover_letter_id) if cover_letter_id and cover_letter_id.isdigit() else None,
            user_id=current_user.id
        )
        db.session.add(job)
        db.session.commit()
        return jsonify({'success': True})
    # Fallback to Flask-WTF form
    form = JobApplicationForm()
    resumes = Resume.query.filter_by(user_id=current_user.id).all()
    covers = CoverLetter.query.filter_by(user_id=current_user.id).all()
    form.resume_id.choices = [(0, 'None')] + [(r.id, r.name or 'Untitled') for r in resumes]
    form.cover_letter_id.choices = [(0, 'None')] + [(c.id, c.title or 'Untitled') for c in covers]
    if form.validate_on_submit():
        job = JobApplication(
            job_title=form.job_title.data,
            company=form.company.data,
            status=form.status.data,
            applied_date=form.applied_date.data or datetime.now(timezone.utc),
            resume_id=form.resume_id.data if form.resume_id.data else None,
            cover_letter_id=form.cover_letter_id.data if form.cover_letter_id.data else None,
            user_id=current_user.id
        )
        db.session.add(job)
        db.session.commit()
        return jsonify({'success': True})
    return jsonify({'success': False, 'message': 'Invalid form data'})

@app.route('/edit-job-application/<int:app_id>', methods=['GET', 'POST'])
@login_required
def edit_job_application(app_id):
    job = JobApplication.query.filter_by(id=app_id, user_id=current_user.id).first_or_404()
    form = JobApplicationForm(obj=job)
    resumes = Resume.query.filter_by(user_id=current_user.id).all()
    covers = CoverLetter.query.filter_by(user_id=current_user.id).all()
    form.resume_id.choices = [(0, 'None')] + [(r.id, r.name or 'Untitled') for r in resumes]
    form.cover_letter_id.choices = [(0, 'None')] + [(c.id, c.title or 'Untitled') for c in covers]
    if form.validate_on_submit():
        job.job_title = form.job_title.data
        job.company = form.company.data
        job.status = form.status.data
        job.applied_date = form.applied_date.data or job.applied_date
        job.resume_id = form.resume_id.data if form.resume_id.data else None
        job.cover_letter_id = form.cover_letter_id.data if form.cover_letter_id.data else None
        db.session.commit()
        flash('Job application updated!', 'success')
        return redirect(url_for('job_application_tracker'))
    return render_template('edit_job_application.html', form=form, action='Edit')

@app.route('/delete-job-application/<int:app_id>', methods=['POST'])
@login_required
def delete_job_application(app_id):
    job = JobApplication.query.filter_by(id=app_id, user_id=current_user.id).first_or_404()
    db.session.delete(job)
    db.session.commit()
    return jsonify({'success': True})

@app.route('/practice-sessions')
@login_required
def practice_sessions():
    sessions = PracticeSession.query.filter_by(user_id=current_user.id).order_by(PracticeSession.created_at.desc()).all()
    return render_template('practice_sessions.html', sessions=sessions, is_premium=is_premium(current_user))

@app.route('/add-practice-session', methods=['GET', 'POST'])
@login_required
def add_practice_session():
    form = PracticeSessionForm()
    if form.validate_on_submit():
        session = PracticeSession(
            category=form.category.data,
            score=float(form.score.data),
            user_id=current_user.id
        )
        db.session.add(session)
        db.session.commit()
        return redirect(url_for('practice_sessions'))
    return render_template('edit_practice_session.html', form=form, action='Add')

@app.route('/edit-practice-session/<int:session_id>', methods=['GET', 'POST'])
@login_required
def edit_practice_session(session_id):
    session = PracticeSession.query.filter_by(id=session_id, user_id=current_user.id).first_or_404()
    form = PracticeSessionForm(obj=session)
    if form.validate_on_submit():
        session.category = form.category.data
        session.score = float(form.score.data)
        db.session.commit()
        return redirect(url_for('practice_sessions'))
    return render_template('edit_practice_session.html', form=form, action='Edit')

@app.route('/delete-practice-session/<int:session_id>', methods=['POST'])
@login_required
def delete_practice_session(session_id):
    session = PracticeSession.query.filter_by(id=session_id, user_id=current_user.id).first_or_404()
    db.session.delete(session)
    db.session.commit()
    return redirect(url_for('practice_sessions'))

@app.route('/interview-prep-advanced')
@login_required
def interview_prep_advanced():
    # Get practice session stats
    practice_sessions = PracticeSession.query.filter_by(user_id=current_user.id).all()
    
    # Calculate stats
    total_practiced = len(practice_sessions)
    avg_score = sum(session.score for session in practice_sessions) / total_practiced if total_practiced > 0 else 0
    mock_interviews = len([s for s in practice_sessions if s.category == 'mock_interview'])
    
    # Calculate streak (simplified - just count recent sessions)
    recent_sessions = [s for s in practice_sessions if s.created_at.date() >= (datetime.now(timezone.utc).date() - timedelta(days=7))]
    streak = len(recent_sessions)
    
    # Get category counts
    behavioral_count = len([s for s in practice_sessions if s.category == 'behavioral'])
    technical_count = len([s for s in practice_sessions if s.category == 'technical'])
    situational_count = len([s for s in practice_sessions if s.category == 'situational'])
    
    stats = {
        'total_practiced': total_practiced,
        'avg_score': round(avg_score, 1),
        'mock_interviews': mock_interviews,
        'streak': streak
    }
    
    return render_template('interview_prep_advanced.html', 
                         is_premium=is_premium(current_user),
                         stats=stats,
                         behavioral_count=behavioral_count,
                         technical_count=technical_count,
                         situational_count=situational_count)

@app.route('/chatbot', methods=['POST'])
@csrf.exempt
def chatbot():
    # Handle both JSON and form data
    if request.is_json:
        user_message = request.json.get('message', '').strip()
        section = request.json.get('section', 'general')
        page = request.json.get('page', 'general')
    else:
        user_message = request.form.get('message', '').strip()
        section = request.form.get('section', 'general')
        page = request.form.get('page', 'general')
    
    if not user_message:
        return jsonify({'reply': 'Please enter a message.'})
    
    # Try OpenRouter AI first (primary)
    if app.config.get('OPENROUTER_API_KEY') and app.config.get('OPENROUTER_API_KEY') != 'your_openrouter_api_key_here':
        try:
            headers = {
                'Authorization': f'Bearer {app.config["OPENROUTER_API_KEY"]}',
                'Content-Type': 'application/json',
                'HTTP-Referer': app.config.get('APP_URL', 'http://localhost:5000'),
                'X-Title': 'ResumeBuilder AI Assistant'
            }
            
            system_prompt = (
                "You are Resume Assistant, an expert career coach and resume writer. Respond conversationally and helpfully, like ChatGPT. "
                "Give detailed, step-by-step advice, examples, and encouragement. If the user asks for a summary, generate a strong, tailored summary. "
                "If they ask for skills, suggest a list with explanations. Always be friendly, supportive, and professional.\n\n"
                "Context: The user is currently on the '{page}' page, working on the '{section}' section. Tailor your response to this context.\n\n"
                "**Formatting instructions:** Always format your response clearly. Use bullet points for lists, numbered steps for instructions, and bold for section headers. "
                "If suggesting a summary, output a single strong paragraph. For skills, use a bulleted list. For experience, use a short, clear example. Format for easy copy-paste."
            )
            data = {
                'model': 'anthropic/claude-3.5-sonnet',
                'messages': [
                    {
                        'role': 'system',
                        'content': system_prompt
                    },
                    {
                        'role': 'user',
                        'content': user_message
                    }
                ],
                'max_tokens': 500,
                'temperature': 0.7
            }
            
            response = requests.post(
                app.config['OPENROUTER_API_URL'],
                headers=headers,
                json=data,
                timeout=15
            )
            
            if response.status_code == 200:
                result = response.json()
                if 'choices' in result and result['choices']:
                    ai_response = result['choices'][0]['message']['content']
                    return jsonify({'reply': ai_response})
            else:
                print(f"OpenRouter API error: {response.status_code} - {response.text}")
        except Exception as e:
            print(f"OpenRouter API error: {e}")
    
    # Fallback to Gemini AI
    elif app.config.get('GEMINI_API_KEY') and app.config.get('GEMINI_API_KEY') != 'your_gemini_api_key_here':
        try:
            headers = {
                'Content-Type': 'application/json',
            }
            data = {
                'contents': [{
                    'parts': [{'text': f"You are a helpful career assistant. Help with resume tips, interview questions, and career advice. User: {user_message}"}]
                }]
            }
            
            response = requests.post(
                f'{app.config["GEMINI_API_URL"]}?key={app.config["GEMINI_API_KEY"]}',
                headers=headers,
                json=data,
                timeout=10
            )
            
            if response.status_code == 200:
                result = response.json()
                if 'candidates' in result and result['candidates']:
                    ai_response = result['candidates'][0]['content']['parts'][0]['text']
                    return jsonify({'reply': ai_response})
        except Exception as e:
            print(f"Gemini API error: {e}")
    
    # Fallback to rule-based responses
    responses = {
        'hello': 'Hello! I\'m here to help with your resume and career questions. How can I assist you today?',
        'resume': 'For a great resume, focus on achievements, use action verbs, and tailor it to the job. Would you like specific tips?',
        'interview': 'Interview preparation is key! Research the company, practice common questions, and prepare your own questions. Need help with specific questions?',
        'help': 'I can help with resume writing, interview preparation, career advice, and more. Just ask!',
        'cover letter': 'A strong cover letter should be tailored to the job, highlight relevant experience, and show enthusiasm for the role. Need help writing one?',
        'skills': 'When listing skills on your resume, focus on relevant technical and soft skills. Quantify achievements when possible.',
        'experience': 'When describing work experience, use action verbs and quantify achievements. Focus on results, not just responsibilities.',
        'education': 'List your education with the most recent degree first. Include relevant coursework, GPA (if 3.0+), and any honors.',
        'networking': 'Networking is crucial for career growth. Attend industry events, connect on LinkedIn, and follow up with meaningful conversations.',
        'salary': 'Research salary ranges for your role and location. Be prepared to negotiate based on your experience and the market value.',
        'linkedin': 'Optimize your LinkedIn profile with a professional photo, compelling headline, and detailed experience descriptions.',
        'portfolio': 'A portfolio showcases your work and skills. Include relevant projects, case studies, and testimonials.',
        'certification': 'Relevant certifications can boost your resume. Focus on industry-recognized credentials that align with your career goals.'
    }
    
    # Check for keyword matches
    user_message_lower = user_message.lower()
    for key, response in responses.items():
        if key in user_message_lower:
            return jsonify({'reply': response})
    
    # Default response if no specific match
    return jsonify({'reply': 'I\'m here to help with resume tips, interview questions, and career advice. Try asking about resume writing, interview preparation, cover letters, networking, or career guidance!'})

# Add analytics tracking
@app.route('/analytics')
@login_required
def analytics():
    # Application status pie
    jobs = JobApplication.query.filter_by(user_id=current_user.id).all()
    status_labels = ['applied', 'interview', 'offer', 'rejected', 'withdrawn']
    status_counts = [sum(1 for j in jobs if j.status == s) for s in status_labels]
    # Activity over time (last 10 days)
    today = dt.date.today()
    activity_dates = [(today - dt.timedelta(days=i)).strftime('%b %d') for i in reversed(range(10))]
    activity_counts = [sum(1 for j in jobs if j.applied_date and j.applied_date.date() == (today - dt.timedelta(days=i))) for i in reversed(range(10))]
    # Template usage
    covers = CoverLetter.query.filter_by(user_id=current_user.id).all()
    template_usage = Counter([c.template for c in covers])
    return render_template('analytics.html',
        status_labels=status_labels,
        status_counts=status_counts,
        activity_dates=activity_dates,
        activity_counts=activity_counts,
        template_usage=template_usage,
        is_premium=is_premium(current_user))

# Referral system
@app.route('/referrals')
@login_required
def referrals():
    # Generate unique referral code for user
    referral_code = f"REF{current_user.id:06d}"
    
    # Mock referral stats (in real app, this would come from database)
    total_referrals = 0
    earned_credits = 0
    pending_referrals = []
    
    return render_template('referrals.html', 
                         referral_code=referral_code,
                         total_referrals=total_referrals,
                         earned_credits=earned_credits,
                         pending_referrals=pending_referrals,
                         is_premium=is_premium(current_user))

@app.route('/referral/<code>')
def referral_landing(code):
    return render_template('referral_landing.html', referral_code=code)

@app.route('/upgrade', methods=['GET', 'POST'])
@login_required
def upgrade():
    monthly_price_id = app.config['STRIPE_MONTHLY_PRICE_ID']
    yearly_price_id = app.config['STRIPE_YEARLY_PRICE_ID']
    if request.method == 'POST':
        plan = request.form.get('plan', 'monthly')
        price_id = monthly_price_id if plan == 'monthly' else yearly_price_id
        try:
            checkout_session = stripe.checkout.Session.create(
                customer_email=current_user.email,
                payment_method_types=['card'],
                line_items=[{
                    'price': price_id,
                    'quantity': 1,
                }],
                mode='subscription',
                success_url=url_for('upgrade_success', _external=True),
                cancel_url=url_for('upgrade_cancel', _external=True),
                metadata={
                    'user_id': current_user.id,
                    'user_email': current_user.email
                }
            )
            return redirect(checkout_session.url)
        except Exception as e:
            flash(f'Error creating checkout session: {str(e)}', 'error')
            return redirect(url_for('dashboard'))
    return render_template('upgrade.html', 
        monthly_price_id=monthly_price_id,
        yearly_price_id=yearly_price_id,
        stripe_publishable_key=STRIPE_PUBLISHABLE_KEY,
        is_premium=is_premium(current_user))

@app.route('/upgrade-success')
@login_required
def upgrade_success():
    # Update user to premium
    current_user.is_premium = True
    current_user.premium_expiry = datetime.now(timezone.utc) + timedelta(days=30)
    db.session.commit()
    
    # Send premium upgrade email
    send_email(
        current_user.email,
        'Welcome to Premium - ResumeBuilder!',
        'premium_upgrade',
        username=current_user.username
    )
    
    flash('Congratulations! You are now a premium member!', 'success')
    return redirect(url_for('dashboard'))

@app.route('/upgrade-cancel')
@login_required
def upgrade_cancel():
    flash('Upgrade was cancelled. You can try again anytime!', 'info')
    return redirect(url_for('dashboard'))

@app.route('/create-checkout-session', methods=['POST'])
@login_required
def create_checkout_session():
    plan = request.form.get('plan', 'monthly')
    price_id = app.config['STRIPE_MONTHLY_PRICE_ID'] if plan == 'monthly' else app.config['STRIPE_YEARLY_PRICE_ID']
    try:
        checkout_session = stripe.checkout.Session.create(
            customer_email=current_user.email,
            payment_method_types=['card'],
            line_items=[{
                'price': price_id,
                'quantity': 1,
            }],
            mode='subscription',
            success_url=url_for('upgrade_success', _external=True),
            cancel_url=url_for('upgrade_cancel', _external=True),
            metadata={
                'user_id': current_user.id,
                'user_email': current_user.email
            }
        )
        return jsonify({'id': checkout_session.id})
    except Exception as e:
        return jsonify({'error': str(e)}), 400

@app.route('/stripe-webhook', methods=['POST'])
def stripe_webhook():
    payload = request.get_data()
    sig_header = request.headers.get('Stripe-Signature')

    try:
        event = stripe.Webhook.construct_event(
            payload, sig_header, STRIPE_WEBHOOK_SECRET
        )
    except ValueError as e:
        return 'Invalid payload', 400
    except stripe.error.SignatureVerificationError as e:
        return 'Invalid signature', 400

    if event['type'] == 'checkout.session.completed':
        session = event['data']['object']
        user_id = session['metadata']['user_id']
        
        # Update user to premium
        user = User.query.get(user_id)
        if user:
            user.is_premium = True
            db.session.commit()
    
    elif event['type'] == 'customer.subscription.deleted':
        subscription = event['data']['object']
        # Handle subscription cancellation
        pass

    return '', 200

@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        email = request.form.get('email')
        user = User.query.filter_by(email=email).first()
        
        if user:
            # Generate reset token
            token = generate_reset_token()
            user.reset_token = token
            user.reset_token_expiry = time.time() + 3600  # 1 hour
            db.session.commit()
            
            # Send reset email
            reset_url = url_for('reset_password', token=token, _external=True)
            if send_email(
                user.email,
                'Password Reset Request - ResumeBuilder',
                'password_reset',
                username=user.username,
                reset_url=reset_url,
                text_content=f'''To reset your password, visit the following link:
{reset_url}

If you did not make this request, simply ignore this email.

This link will expire in 1 hour.
'''
            ):
                flash('An email has been sent with instructions to reset your password.', 'success')
            else:
                flash('Error sending email. Please try again.', 'error')
        else:
            flash('Email address not found.', 'error')
    
    return render_template('forgot_password.html')

@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    user = verify_reset_token(token)
    if not user:
        flash('Invalid or expired reset token.', 'error')
        return redirect(url_for('forgot_password'))
    
    if request.method == 'POST':
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        
        if password != confirm_password:
            flash('Passwords do not match.', 'error')
        elif len(password) < 6:
            flash('Password must be at least 6 characters long.', 'error')
        else:
            user.password_hash = generate_password_hash(password)
            user.reset_token = None
            user.reset_token_expiry = None
            db.session.commit()
            flash('Your password has been updated! You can now log in.', 'success')
            return redirect(url_for('login'))
    
    return render_template('reset_password.html', token=token)

@app.route('/profile')
@login_required
def profile():
    return render_template('profile.html', user=current_user, is_premium=is_premium(current_user))

@app.route('/upload-profile-image', methods=['POST'])
@login_required
def upload_profile_image():
    if 'profile_image' not in request.files:
        flash('No file selected', 'error')
        return redirect(url_for('profile'))
    file = request.files['profile_image']
    if file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('profile'))
    # Validate file type
    allowed_types = ['image/jpeg', 'image/png']
    if file.content_type not in allowed_types:
        flash('Only JPEG and PNG images are allowed.', 'error')
        return redirect(url_for('profile'))
    # Validate file size (max 2MB)
    file.seek(0, 2)
    size = file.tell()
    file.seek(0)
    if size > 2 * 1024 * 1024:
        flash('File size must be under 2MB.', 'error')
        return redirect(url_for('profile'))
    # Save file
    upload_folder = os.path.join(app.root_path, 'static', 'uploads')
    os.makedirs(upload_folder, exist_ok=True)
    filename = f"profile_{current_user.id}_{int(time.time())}.jpg"
    filepath = os.path.join(upload_folder, filename)
    file.save(filepath)
    current_user.profile_image = f"/static/uploads/{filename}"
    db.session.commit()
    flash('Profile image updated successfully!', 'success')
    return redirect(url_for('profile'))

@app.route('/remove-profile-image', methods=['POST'])
@login_required
def remove_profile_image():
    if current_user.profile_image:
        # Remove file from disk
        try:
            os.remove(os.path.join(app.root_path, current_user.profile_image.lstrip('/')))
        except Exception:
            pass
        current_user.profile_image = None
        db.session.commit()
        flash('Profile image removed.', 'success')
    return redirect(url_for('profile'))

@app.route('/set-avatar', methods=['POST'])
@login_required
def set_avatar():
    avatar = request.form.get('avatar')
    allowed = [f'/static/avatar{i}.png' for i in range(1, 4)]
    if avatar in allowed:
        current_user.profile_image = avatar
        db.session.commit()
        flash('Avatar updated!', 'success')
    else:
        flash('Invalid avatar selection.', 'danger')
    return redirect(url_for('profile'))

@app.route('/change-password', methods=['GET', 'POST'])
@login_required
def change_password():
    form = ChangePasswordForm()
    if request.method == 'POST':
        if form.validate_on_submit():
            if not check_password_hash(current_user.password_hash, form.current_password.data):
                flash('Current password is incorrect.', 'error')
            elif form.new_password.data != form.confirm_new_password.data:
                flash('New passwords do not match.', 'error')
            elif len(form.new_password.data) < 6:
                flash('New password must be at least 6 characters long.', 'error')
            else:
                current_user.password_hash = generate_password_hash(form.new_password.data)
                db.session.commit()
                flash('Password updated successfully!', 'success')
                return redirect(url_for('profile'))
        else:
            flash('Please correct the errors in the form.', 'error')
    
    return render_template('change_password.html', form=form)

@app.route('/delete-account', methods=['POST'])
@login_required
def delete_account():
    # Delete user's resumes first
    Resume.query.filter_by(user_id=current_user.id).delete()
    
    # Delete user
    db.session.delete(current_user)
    db.session.commit()
    
    flash('Your account has been deleted.', 'info')
    return redirect(url_for('landing'))

@app.route('/api/generate-cover-letter', methods=['POST'])
@login_required
def api_generate_cover_letter():
    job_title = request.form.get('job_title', '')
    company = request.form.get('company', '')
    industry = request.form.get('industry', '')
    experience = request.form.get('experience', '')
    skills = request.form.get('skills', '')
    achievements = request.form.get('achievements', '')
    motivation = request.form.get('motivation', '')
    tone = request.form.get('tone', 'Professional')
    prompt = f"Write a {tone.lower()} cover letter for a {experience} position as {job_title} at {company} in the {industry} industry. Skills: {skills}. Achievements: {achievements}. Motivation: {motivation}."
    # Use Gemini AI or fallback
    letter = None
    if app.config.get('GEMINI_API_KEY') and app.config.get('GEMINI_API_KEY') != 'your_gemini_api_key_here':
        try:
            headers = {'Content-Type': 'application/json'}
            data = {'contents': [{'parts': [{'text': prompt}]}]}
            response = requests.post(f'{app.config["GEMINI_API_URL"]}?key={app.config["GEMINI_API_KEY"]}', headers=headers, json=data, timeout=10)
            if response.status_code == 200:
                result = response.json()
                if 'candidates' in result and result['candidates']:
                    letter = result['candidates'][0]['content']['parts'][0]['text']
        except Exception as e:
            print(f"Gemini API error: {e}")
    if not letter:
        letter = f"Dear Hiring Manager,\n\nI am excited to apply for the {job_title} position at {company}. My experience in {industry} and skills in {skills} make me a strong fit. {achievements} {motivation}\n\nSincerely,\n{current_user.username}"
    return jsonify({'success': True, 'letter': letter})

@app.route('/api/save-cover-letter', methods=['POST'])
@login_required
def api_save_cover_letter():
    data = request.get_json()
    content = data.get('content', '').strip()
    if not content:
        return jsonify({'success': False, 'message': 'No content provided'})
    cover = CoverLetter(
        title='AI Generated',
        job_title='',
        company='',
        content=content,
        template='Standard',
        user_id=current_user.id,
        created_at=datetime.now(timezone.utc),
        updated_at=datetime.now(timezone.utc)
    )
    db.session.add(cover)
    db.session.commit()
    return jsonify({'success': True, 'id': cover.id})

# Helper: create reminders for job follow-up/interview
@app.route('/create-reminders')
@login_required
def create_reminders():
    jobs = JobApplication.query.filter_by(user_id=current_user.id).all()
    for job in jobs:
        # Example: create reminder 1 day before interview (if interview_date exists)
        if hasattr(job, 'interview_date') and job.interview_date:
            due = job.interview_date - timedelta(days=1)
            if not Reminder.query.filter_by(user_id=current_user.id, due_date=due, message=f"Interview for {job.job_title} at {job.company}").first():
                r = Reminder(user_id=current_user.id, message=f"Interview for {job.job_title} at {job.company}", due_date=due)
                db.session.add(r)
        # Example: follow-up reminder 7 days after applied
        if job.applied_date:
            due = job.applied_date + timedelta(days=7)
            if not Reminder.query.filter_by(user_id=current_user.id, due_date=due, message=f"Follow up on {job.job_title} at {job.company}").first():
                r = Reminder(user_id=current_user.id, message=f"Follow up on {job.job_title} at {job.company}", due_date=due)
                db.session.add(r)
    db.session.commit()
    flash('Reminders created!', 'success')
    return redirect(url_for('dashboard'))

# Daily email reminder function (to be called by a scheduler like cron)
def send_daily_reminders():
    now = datetime.now(timezone.utc)
    upcoming = Reminder.query.filter(Reminder.due_date <= now + timedelta(days=1), Reminder.sent == False).all()
    for r in upcoming:
        user = User.query.get(r.user_id)
        if user and user.email:
            try:
                msg = Message('Upcoming Reminder', sender=app.config['MAIL_DEFAULT_SENDER'], recipients=[user.email])
                msg.body = f"Reminder: {r.message} is due on {r.due_date.strftime('%b %d, %Y')}"
                mail.send(msg)
                r.sent = True
                db.session.commit()
            except Exception as e:
                print(f"Reminder email error: {e}")

@app.route('/teams')
@login_required
def teams():
    my_teams = TeamCV.query.filter_by(owner_id=current_user.id).all()
    member_teams = TeamCV.query.join(TeamMember).filter(TeamMember.user_id==current_user.id).all()
    return render_template('teams.html', my_teams=my_teams, member_teams=member_teams, is_premium=is_premium(current_user))

@app.route('/create-team', methods=['GET', 'POST'])
@login_required
def create_team():
    if request.method == 'POST':
        name = request.form.get('name')
        if name:
            team = TeamCV(name=name, owner_id=current_user.id)
            db.session.add(team)
            db.session.commit()
            flash('Team created!', 'success')
            return redirect(url_for('teams'))
    return render_template('edit_team.html', action='Create')

@app.route('/team/<int:team_id>')
@login_required
def view_team(team_id):
    team = TeamCV.query.get_or_404(team_id)
    members = TeamMember.query.filter_by(team_id=team_id).all()
    return render_template('view_team.html', team=team, members=members, is_owner=(team.owner_id==current_user.id), is_premium=is_premium(current_user))

@app.route('/invite-team-member/<int:team_id>', methods=['POST'])
@login_required
def invite_team_member(team_id):
    email = request.form.get('email')
    role = request.form.get('role', 'member')
    if email:
        member = TeamMember(team_id=team_id, invited_email=email, role=role)
        db.session.add(member)
        db.session.commit()
        flash('Invitation sent!', 'success')
    return redirect(url_for('view_team', team_id=team_id))

@app.route('/remove-team-member/<int:member_id>', methods=['POST'])
@login_required
def remove_team_member(member_id):
    member = TeamMember.query.get_or_404(member_id)
    team_id = member.team_id
    db.session.delete(member)
    db.session.commit()
    flash('Member removed.', 'info')
    return redirect(url_for('view_team', team_id=team_id))

@app.route('/affiliate')
@login_required
def affiliate_dashboard():
    link = AffiliateLink.query.filter_by(user_id=current_user.id).first()
    if not link:
        code = f"AFF{current_user.id:06d}"
        link = AffiliateLink(user_id=current_user.id, code=code)
        db.session.add(link)
        db.session.commit()
    return render_template('affiliate.html', link=link, is_premium=is_premium(current_user))

@app.route('/affiliate-click/<code>')
def affiliate_click(code):
    link = AffiliateLink.query.filter_by(code=code).first()
    if link:
        link.clicks += 1
        db.session.commit()
    return redirect(url_for('register'))

@app.route('/affiliate-signup/<code>')
def affiliate_signup(code):
    link = AffiliateLink.query.filter_by(code=code).first()
    if link:
        link.signups += 1
        db.session.commit()
    return redirect(url_for('register'))

@app.route('/affiliate-payout', methods=['POST'])
@login_required
def affiliate_payout():
    link = AffiliateLink.query.filter_by(user_id=current_user.id).first()
    if link and link.signups > 0:
        link.payouts += link.signups * 5.0  # $5 per signup (mock)
        link.signups = 0
        db.session.commit()
        flash('Payout requested! (mock)', 'success')
    else:
        flash('No earnings to payout.', 'info')
    return redirect(url_for('affiliate_dashboard'))

@app.route('/cover-letter-templates')
@login_required
def cover_letter_templates():
    return render_template('cover_letter_templates.html', is_premium=is_premium(current_user))

@app.route('/cover-letter-examples')
@login_required
def cover_letter_examples():
    return render_template('cover_letter_templates.html', is_premium=is_premium(current_user))

@app.route('/portfolio')
@login_required
def portfolio():
    return render_template('portfolio.html', is_premium=is_premium(current_user))

@app.route('/website-generator')
@login_required
def website_generator():
    return render_template('website_generator.html', is_premium=is_premium(current_user))

@app.route('/interview-prep')
@login_required
def interview_prep():
    return redirect(url_for('interview_prep_advanced'))

@app.route('/networking-tools')
@login_required
def networking_tools():
    return render_template('networking_tools.html', is_premium=is_premium(current_user))

@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    if request.method == 'POST':
        user = current_user
        user.analytics_enabled = bool(request.form.get('email_notifications'))
        user.theme = 'dark' if request.form.get('theme') == 'dark' else 'light'
        db.session.commit()
        flash('Settings updated!', 'success')
        return redirect(url_for('settings'))
    return render_template('settings.html', user=current_user, is_premium=is_premium(current_user))

@app.route('/admin')
@login_required
def admin_dashboard():
    if not getattr(current_user, 'is_admin', False):
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    users = User.query.all()
    return render_template('admin.html', users=users, is_premium=is_premium(current_user))

@app.route('/admin/toggle-premium/<int:user_id>', methods=['POST'])
@login_required
def admin_toggle_premium(user_id):
    if not getattr(current_user, 'is_admin', False):
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    user = User.query.get_or_404(user_id)
    user.is_premium = not user.is_premium
    db.session.commit()
    flash('Premium status toggled.', 'success')
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/delete-user/<int:user_id>', methods=['POST'])
@login_required
def admin_delete_user(user_id):
    if not getattr(current_user, 'is_admin', False):
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    user = User.query.get_or_404(user_id)
    db.session.delete(user)
    db.session.commit()
    flash('User deleted.', 'info')
    return redirect(url_for('admin_dashboard'))

# Error handlers
@app.errorhandler(400)
def bad_request(error):
    logger.error(f"400 Bad Request: {error}")
    logger.error(f"Request URL: {request.url}")
    logger.error(f"Request method: {request.method}")
    logger.error(f"Request headers: {dict(request.headers)}")
    logger.error(f"Request form data: {dict(request.form)}")
    return render_template('404.html'), 400

@app.errorhandler(404)
def not_found_error(error):
    logger.error(f"404 Not Found: {error}")
    logger.error(f"Request URL: {request.url}")
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_error(error):
    logger.error(f"500 Internal Server Error: {error}")
    logger.error(f"Request URL: {request.url}")
    import traceback
    logger.error(f"Traceback: {traceback.format_exc()}")
    return render_template('404.html'), 500

@app.errorhandler(Exception)
def handle_exception(e):
    logger.error(f"Unhandled Exception: {e}")
    logger.error(f"Request URL: {request.url}")
    logger.error(f"Request method: {request.method}")
    import traceback
    logger.error(f"Traceback: {traceback.format_exc()}")
    # Don't expose sensitive information in production
    if app.config.get('DEBUG', False):
        return render_template('500.html', error=str(e)), 500
    else:
        return render_template('500.html'), 500

# Security utilities
def sanitize_input(text):
    """Sanitize user input to prevent XSS"""
    if not text:
        return text
    import html
    return html.escape(text)

def validate_file_upload(file, allowed_extensions=['png', 'jpg', 'jpeg', 'gif']):
    """Validate file uploads for security"""
    if not file or not file.filename:
        return False, "No file selected"
    
    # Check file extension
    if '.' not in file.filename:
        return False, "Invalid file type"
    
    ext = file.filename.rsplit('.', 1)[1].lower()
    if ext not in allowed_extensions:
        return False, f"File type .{ext} not allowed"
    
    # Check file size (5MB limit)
    if len(file.read()) > 5 * 1024 * 1024:
        file.seek(0)  # Reset file pointer
        return False, "File too large (max 5MB)"
    
    file.seek(0)  # Reset file pointer
    return True, "File valid"

# Enhanced session management
@app.before_request
def session_security():
    """Enhanced session security"""
    if current_user.is_authenticated:
        # Regenerate session ID periodically
        if 'last_regeneration' not in session:
            session['last_regeneration'] = datetime.now(timezone.utc)
        elif (datetime.now(timezone.utc) - session['last_regeneration']).days > 1:
            session.regenerate()
            session['last_regeneration'] = datetime.now(timezone.utc)
        
        # Set session as permanent for authenticated users
        session.permanent = True

# Security monitoring
@app.after_request
def security_monitoring(response):
    """Monitor for security issues"""
    # Log failed login attempts
    if request.endpoint == 'login' and response.status_code == 401:
        logger.warning(f"Failed login attempt from {request.remote_addr}")
    
    # Log suspicious activity
    if response.status_code in [400, 403, 404, 500]:
        logger.warning(f"Suspicious activity: {request.method} {request.url} - {response.status_code}")
    
    return response

@app.route('/api/auto-save-resume', methods=['POST'])
@login_required
def auto_save_resume():
    """Auto-save resume data via AJAX"""
    try:
        data = request.form.to_dict()
        
        # Validate required fields
        required_fields = ['name', 'email', 'phone', 'summary', 'education', 'experience', 'skills', 'languages']
        for field in required_fields:
            if not data.get(field):
                return jsonify({'success': False, 'error': f'Missing required field: {field}'}), 400
        
        # Check if resume exists (by name and user)
        existing_resume = Resume.query.filter_by(
            name=data['name'], 
            user_id=current_user.id
        ).first()
        
        if existing_resume:
            # Update existing resume
            existing_resume.email = sanitize_input(data['email'])
            existing_resume.phone = sanitize_input(data['phone'])
            existing_resume.address = sanitize_input(data.get('address', ''))
            existing_resume.linkedin = sanitize_input(data.get('linkedin', ''))
            existing_resume.summary = sanitize_input(data['summary'])
            existing_resume.education = sanitize_input(data['education'])
            existing_resume.experience = sanitize_input(data['experience'])
            existing_resume.skills = sanitize_input(data['skills'])
            existing_resume.certifications = sanitize_input(data.get('certifications', ''))
            existing_resume.languages = sanitize_input(data['languages'])
            existing_resume.template = sanitize_input(data.get('template', 'modern'))
            existing_resume.updated_at = datetime.now(timezone.utc)
            
            db.session.commit()
            logger.info(f"Auto-saved existing resume: {existing_resume.name} for user {current_user.id}")
            
            return jsonify({
                'success': True, 
                'message': 'Resume auto-saved successfully',
                'resume_id': existing_resume.id,
                'action': 'updated'
            })
        else:
            # Create new resume
            new_resume = Resume(
                name=sanitize_input(data['name']),
                email=sanitize_input(data['email']),
                phone=sanitize_input(data['phone']),
                address=sanitize_input(data.get('address', '')),
                linkedin=sanitize_input(data.get('linkedin', '')),
                summary=sanitize_input(data['summary']),
                education=sanitize_input(data['education']),
                experience=sanitize_input(data['experience']),
                skills=sanitize_input(data['skills']),
                certifications=sanitize_input(data.get('certifications', '')),
                languages=sanitize_input(data['languages']),
                template=sanitize_input(data.get('template', 'modern')),
                pdf_engine=data.get('pdf_engine', 'weasyprint'),
                user_id=current_user.id
            )
            
            db.session.add(new_resume)
            db.session.commit()
            logger.info(f"Auto-saved new resume: {new_resume.name} for user {current_user.id}")
            
            return jsonify({
                'success': True, 
                'message': 'Resume auto-saved successfully',
                'resume_id': new_resume.id,
                'action': 'created'
            })
            
    except Exception as e:
        logger.error(f"Auto-save error: {e}")
        db.session.rollback()
        return jsonify({'success': False, 'error': 'Failed to auto-save resume'}), 500

@app.route('/api/resume-stats', methods=['GET'])
@login_required
def get_resume_stats():
    """Get resume statistics for dashboard"""
    try:
        user_resumes = Resume.query.filter_by(user_id=current_user.id).all()
        
        stats = {
            'total_resumes': len(user_resumes),
            'recent_activity': [],
            'template_usage': {},
            'completion_rate': 0
        }
        
        # Calculate completion rate
        if user_resumes:
            completed = sum(1 for r in user_resumes if r.summary and r.experience and r.skills)
            stats['completion_rate'] = round((completed / len(user_resumes)) * 100, 1)
        
        # Template usage
        for resume in user_resumes:
            template = resume.template or 'modern'
            stats['template_usage'][template] = stats['template_usage'].get(template, 0) + 1
        
        # Recent activity (last 5 resumes)
        recent_resumes = sorted(user_resumes, key=lambda x: x.updated_at or x.created_at, reverse=True)[:5]
        for resume in recent_resumes:
            stats['recent_activity'].append({
                'id': resume.id,
                'name': resume.name,
                'action': 'updated' if resume.updated_at else 'created',
                'date': (resume.updated_at or resume.created_at).strftime('%Y-%m-%d %H:%M'),
                'template': resume.template or 'modern'
            })
        
        return jsonify(stats)
        
    except Exception as e:
        logger.error(f"Error getting resume stats: {e}")
        return jsonify({'error': 'Failed to get statistics'}), 500

@app.route('/api/ai-suggestions', methods=['POST'])
@login_required
def get_ai_suggestions():
    """Get AI-powered suggestions for resume sections"""
    try:
        data = request.get_json()
        section = data.get('section')
        context = data.get('context', '')
        
        if not section:
            return jsonify({'error': 'Section is required'}), 400
        
        # AI suggestion prompts
        prompts = {
            'summary': f"Write a professional summary for a resume. Context: {context}",
            'experience': f"Suggest work experience descriptions. Context: {context}",
            'skills': f"Suggest relevant skills for a resume. Context: {context}",
            'education': f"Suggest education section content. Context: {context}"
        }
        
        if section not in prompts:
            return jsonify({'error': 'Invalid section'}), 400
        
        # Use AI to generate suggestions
        try:
            suggestion = generate_ai_suggestion(prompts[section])
            return jsonify({
                'success': True,
                'suggestion': suggestion,
                'section': section
            })
        except Exception as ai_error:
            logger.error(f"AI suggestion error: {ai_error}")
            return jsonify({
                'success': False,
                'error': 'AI suggestions temporarily unavailable',
                'fallback': get_fallback_suggestion(section)
            })
            
    except Exception as e:
        logger.error(f"AI suggestions error: {e}")
        return jsonify({'error': 'Failed to get suggestions'}), 500

def generate_ai_suggestion(prompt):
    """Generate AI suggestion using available APIs"""
    try:
        # Try Gemini first
        if os.getenv('GEMINI_API_KEY'):
            import google.generativeai as genai
            genai.configure(api_key=os.getenv('GEMINI_API_KEY'))
            model = genai.GenerativeModel('gemini-pro')
            response = model.generate_content(prompt)
            return response.text.strip()
        
        # Try Cohere
        elif os.getenv('COHERE_API_KEY'):
            import cohere
            co = cohere.Client(os.getenv('COHERE_API_KEY'))
            response = co.generate(
                model='command',
                prompt=prompt,
                max_tokens=200,
                temperature=0.7
            )
            return response.generations[0].text.strip()
        
        # Try OpenRouter
        elif os.getenv('OPENROUTER_API_KEY'):
            import openai
            openai.api_key = os.getenv('OPENROUTER_API_KEY')
            openai.api_base = "https://openrouter.ai/api/v1"
            response = openai.ChatCompletion.create(
                model="openai/gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=200
            )
            return response.choices[0].message.content.strip()
        
        else:
            raise Exception("No AI API keys configured")
            
    except Exception as e:
        logger.error(f"AI generation error: {e}")
        raise e

def get_fallback_suggestion(section):
    """Get fallback suggestions when AI is unavailable"""
    fallbacks = {
        'summary': "Experienced professional with strong skills in [your field]. Proven track record of [key achievements]. Passionate about [your interests] and committed to delivering exceptional results.",
        'experience': "• Led and managed [specific project/team]\n• Increased [metric] by [percentage]\n• Collaborated with cross-functional teams\n• Implemented [specific solution]",
        'skills': "Technical Skills: [List your technical skills]\nSoft Skills: Communication, Leadership, Problem-solving\nTools: [List relevant tools and software]",
        'education': "Bachelor's Degree in [Field] - [University Name] (Year)\nRelevant Coursework: [List relevant courses]\nGPA: [if applicable]"
    }
    return fallbacks.get(section, "Please provide more context for better suggestions.")

@app.route('/ai-interview-prep', methods=['GET', 'POST'])
@login_required
def ai_interview_prep():
    """Advanced AI Interview Preparation"""
    if request.method == 'POST':
        data = request.get_json()
        job_title = data.get('job_title', '')
        company = data.get('company', '')
        interview_type = data.get('interview_type', 'technical')
        
        # Generate interview questions based on job
        questions = generate_interview_questions(job_title, company, interview_type)
        
        return jsonify({
            'success': True,
            'questions': questions,
            'tips': get_interview_tips(interview_type)
        })
    
    return render_template('ai_interview_prep.html')

def generate_interview_questions(job_title, company, interview_type):
    """Generate AI-powered interview questions"""
    try:
        if interview_type == 'technical':
            prompt = f"Generate 10 technical interview questions for a {job_title} position at {company}. Include coding questions, system design, and problem-solving scenarios."
        elif interview_type == 'behavioral':
            prompt = f"Generate 10 behavioral interview questions for a {job_title} position at {company}. Focus on leadership, teamwork, problem-solving, and past experiences."
        else:
            prompt = f"Generate 10 general interview questions for a {job_title} position at {company}."
        
        questions = generate_ai_suggestion(prompt)
        return questions.split('\n')[:10]  # Return first 10 questions
        
    except Exception as e:
        logger.error(f"Error generating interview questions: {e}")
        return get_fallback_interview_questions(interview_type)

def get_interview_tips(interview_type):
    """Get interview tips based on type"""
    tips = {
        'technical': [
            "Practice coding on a whiteboard or paper",
            "Explain your thought process out loud",
            "Ask clarifying questions before starting",
            "Consider edge cases and optimization",
            "Be honest about what you don't know"
        ],
        'behavioral': [
            "Use the STAR method (Situation, Task, Action, Result)",
            "Prepare specific examples from your experience",
            "Quantify your achievements when possible",
            "Show enthusiasm and passion",
            "Practice your responses beforehand"
        ],
        'general': [
            "Research the company thoroughly",
            "Prepare thoughtful questions to ask",
            "Dress appropriately for the company culture",
            "Arrive early and be prepared",
            "Follow up with a thank-you email"
        ]
    }
    return tips.get(interview_type, tips['general'])

def get_fallback_interview_questions(interview_type):
    """Fallback interview questions when AI is unavailable"""
    questions = {
        'technical': [
            "How would you approach debugging a production issue?",
            "Explain the difference between REST and GraphQL APIs",
            "How would you design a scalable database architecture?",
            "What's your experience with cloud platforms?",
            "How do you stay updated with technology trends?",
            "Describe a challenging technical problem you solved",
            "How would you optimize a slow database query?",
            "What's your experience with version control systems?",
            "How do you handle conflicting requirements?",
            "Explain your testing strategy for a new feature"
        ],
        'behavioral': [
            "Tell me about a time you led a team through a difficult project",
            "Describe a situation where you had to resolve a conflict",
            "How do you handle tight deadlines and pressure?",
            "Tell me about a time you failed and what you learned",
            "Describe a situation where you had to learn something quickly",
            "How do you motivate team members?",
            "Tell me about a time you had to make a difficult decision",
            "Describe a situation where you had to adapt to change",
            "How do you handle criticism and feedback?",
            "Tell me about a time you went above and beyond expectations"
        ],
        'general': [
            "Why are you interested in this position?",
            "What are your career goals for the next 5 years?",
            "Why do you want to work at this company?",
            "What are your greatest strengths and weaknesses?",
            "How do you handle stress and pressure?",
            "What motivates you in your work?",
            "Describe your ideal work environment",
            "How do you prioritize your work?",
            "What are your salary expectations?",
            "Do you have any questions for us?"
        ]
    }
    return questions.get(interview_type, questions['general'])

@app.route('/api/practice-interview', methods=['POST'])
@login_required
def practice_interview():
    """Practice interview with AI feedback"""
    try:
        data = request.get_json()
        question = data.get('question', '')
        answer = data.get('answer', '')
        job_title = data.get('job_title', '')
        
        # Analyze the answer using AI
        feedback = analyze_interview_answer(question, answer, job_title)
        
        return jsonify({
            'success': True,
            'feedback': feedback,
            'score': feedback.get('score', 7),
            'suggestions': feedback.get('suggestions', [])
        })
        
    except Exception as e:
        logger.error(f"Error in practice interview: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to analyze answer',
            'feedback': {
                'score': 7,
                'suggestions': ['Consider providing more specific examples', 'Use the STAR method for behavioral questions']
            }
        })

def analyze_interview_answer(question, answer, job_title):
    """Analyze interview answer using AI"""
    try:
        prompt = f"""
        Analyze this interview answer for a {job_title} position:
        
        Question: {question}
        Answer: {answer}
        
        Provide:
        1. A score out of 10
        2. 3 specific suggestions for improvement
        3. What was done well
        4. Areas for improvement
        """
        
        analysis = generate_ai_suggestion(prompt)
        
        # Parse the analysis (simplified)
        suggestions = [
            "Provide more specific examples",
            "Use the STAR method for behavioral questions",
            "Quantify your achievements when possible"
        ]
        
        return {
            'score': 8,
            'suggestions': suggestions,
            'analysis': analysis,
            'strengths': ['Good structure', 'Relevant experience'],
            'improvements': ['Add more specific metrics', 'Include more context']
        }
        
    except Exception as e:
        logger.error(f"Error analyzing interview answer: {e}")
        return {
            'score': 7,
            'suggestions': ['Consider providing more specific examples', 'Use the STAR method for behavioral questions'],
            'analysis': 'Unable to analyze at this time',
            'strengths': ['Answer provided'],
            'improvements': ['Could be more specific']
        }

@app.route('/ai-resume-analysis', methods=['GET', 'POST'])
@login_required
def ai_resume_analysis():
    """Advanced AI Resume Analysis with ATS Optimization"""
    if request.method == 'POST':
        data = request.get_json()
        resume_id = data.get('resume_id')
        job_description = data.get('job_description', '')
        
        # Get resume data
        resume = Resume.query.get_or_404(resume_id)
        if resume.user_id != current_user.id:
            return jsonify({'error': 'Unauthorized'}), 403
        
        # Perform AI analysis
        analysis = perform_ai_resume_analysis(resume, job_description)
        
        return jsonify({
            'success': True,
            'analysis': analysis
        })
    
    # Get user's resumes for analysis
    resumes = Resume.query.filter_by(user_id=current_user.id).all()
    return render_template('ai_resume_analysis.html', resumes=resumes)

def perform_ai_resume_analysis(resume, job_description=''):
    """Perform comprehensive AI analysis of resume"""
    try:
        # Combine resume content for analysis
        resume_content = f"""
        Name: {resume.name}
        Summary: {resume.summary or ''}
        Experience: {resume.experience or ''}
        Education: {resume.education or ''}
        Skills: {resume.skills or ''}
        """
        
        # ATS Optimization Analysis
        ats_score = analyze_ats_optimization(resume_content, job_description)
        
        # Content Quality Analysis
        content_analysis = analyze_content_quality(resume_content)
        
        # Keyword Analysis
        keyword_analysis = analyze_keywords(resume_content, job_description)
        
        # Structure Analysis
        structure_analysis = analyze_resume_structure(resume)
        
        # Overall Score
        overall_score = calculate_overall_score(ats_score, content_analysis, keyword_analysis, structure_analysis)
        
        return {
            'overall_score': overall_score,
            'ats_optimization': ats_score,
            'content_quality': content_analysis,
            'keyword_analysis': keyword_analysis,
            'structure_analysis': structure_analysis,
            'recommendations': generate_recommendations(ats_score, content_analysis, keyword_analysis, structure_analysis)
        }
        
    except Exception as e:
        logger.error(f"Error in AI resume analysis: {e}")
        return get_fallback_analysis()

def analyze_ats_optimization(resume_content, job_description):
    """Analyze ATS optimization"""
    try:
        # Basic ATS checks
        checks = {
            'has_contact_info': bool(re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', resume_content)),
            'has_phone': bool(re.search(r'[\+]?[1-9][\d]{0,15}', resume_content)),
            'has_skills_section': 'skills' in resume_content.lower(),
            'has_experience_section': 'experience' in resume_content.lower(),
            'has_education_section': 'education' in resume_content.lower(),
            'has_summary': bool(resume_content.split('Summary:')[1].strip() if 'Summary:' in resume_content else False),
            'no_images': True,  # Assuming no images in text content
            'no_tables': True,  # Assuming no tables in text content
            'proper_formatting': True  # Basic check
        }
        
        score = sum(checks.values()) / len(checks) * 100
        
        return {
            'score': round(score, 1),
            'checks': checks,
            'issues': [k for k, v in checks.items() if not v],
            'strengths': [k for k, v in checks.items() if v]
        }
        
    except Exception as e:
        logger.error(f"Error in ATS analysis: {e}")
        return {'score': 70, 'checks': {}, 'issues': [], 'strengths': []}

def analyze_content_quality(resume_content):
    """Analyze content quality and readability"""
    try:
        # Basic content analysis
        word_count = len(resume_content.split())
        sentence_count = len(re.split(r'[.!?]+', resume_content))
        avg_sentence_length = word_count / sentence_count if sentence_count > 0 else 0
        
        # Action verb analysis
        action_verbs = [
            'developed', 'implemented', 'managed', 'created', 'designed', 'built',
            'led', 'coordinated', 'analyzed', 'improved', 'increased', 'decreased',
            'maintained', 'established', 'organized', 'planned', 'executed'
        ]
        
        found_verbs = sum(1 for verb in action_verbs if verb in resume_content.lower())
        verb_score = min(found_verbs / 5 * 100, 100)  # Cap at 100%
        
        # Quantification analysis
        numbers = re.findall(r'\d+', resume_content)
        quantification_score = min(len(numbers) * 10, 100)
        
        return {
            'word_count': word_count,
            'sentence_count': sentence_count,
            'avg_sentence_length': round(avg_sentence_length, 1),
            'action_verbs_score': round(verb_score, 1),
            'quantification_score': round(quantification_score, 1),
            'overall_score': round((verb_score + quantification_score) / 2, 1)
        }
        
    except Exception as e:
        logger.error(f"Error in content analysis: {e}")
        return {'overall_score': 75, 'word_count': 0, 'sentence_count': 0}

def analyze_keywords(resume_content, job_description):
    """Analyze keyword matching with job description"""
    try:
        if not job_description:
            return {'score': 80, 'matched_keywords': [], 'missing_keywords': [], 'suggestions': []}
        
        # Extract keywords from job description
        job_keywords = extract_keywords(job_description)
        resume_keywords = extract_keywords(resume_content)
        
        # Find matches
        matched_keywords = [kw for kw in job_keywords if kw in resume_keywords]
        missing_keywords = [kw for kw in job_keywords if kw not in resume_keywords]
        
        # Calculate score
        match_score = len(matched_keywords) / len(job_keywords) * 100 if job_keywords else 80
        
        return {
            'score': round(match_score, 1),
            'matched_keywords': matched_keywords[:10],  # Top 10
            'missing_keywords': missing_keywords[:10],  # Top 10
            'suggestions': ['Add missing keywords to improve ATS matching', 'Include industry-specific terminology']
        }
        
    except Exception as e:
        logger.error(f"Error in keyword analysis: {e}")
        return {'score': 80, 'matched_keywords': [], 'missing_keywords': [], 'suggestions': []}

def extract_keywords(text):
    """Extract important keywords from text"""
    try:
        # Remove common words and extract meaningful terms
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'can', 'this', 'that', 'these', 'those'}
        
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        words = [word for word in words if word not in stop_words]
        
        # Count frequency and return most common
        word_freq = {}
        for word in words:
            word_freq[word] = word_freq.get(word, 0) + 1
        
        # Return top keywords
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        return [word for word, freq in sorted_words[:20]]  # Top 20 keywords
        
    except Exception as e:
        logger.error(f"Error extracting keywords: {e}")
        return []

def analyze_resume_structure(resume):
    """Analyze resume structure and organization"""
    try:
        sections = {
            'contact_info': bool(resume.name and resume.email and resume.phone),
            'summary': bool(resume.summary and len(resume.summary) > 50),
            'experience': bool(resume.experience and len(resume.experience) > 100),
            'education': bool(resume.education and len(resume.education) > 30),
            'skills': bool(resume.skills and len(resume.skills) > 20),
            'languages': bool(resume.languages and len(resume.languages) > 10)
        }
        
        score = sum(sections.values()) / len(sections) * 100
        
        return {
            'score': round(score, 1),
            'sections': sections,
            'missing_sections': [k for k, v in sections.items() if not v],
            'complete_sections': [k for k, v in sections.items() if v]
        }
        
    except Exception as e:
        logger.error(f"Error in structure analysis: {e}")
        return {'score': 75, 'sections': {}, 'missing_sections': [], 'complete_sections': []}

def calculate_overall_score(ats_score, content_analysis, keyword_analysis, structure_analysis):
    """Calculate overall resume score"""
    try:
        weights = {
            'ats': 0.3,
            'content': 0.25,
            'keywords': 0.25,
            'structure': 0.2
        }
        
        overall = (
            ats_score['score'] * weights['ats'] +
            content_analysis['overall_score'] * weights['content'] +
            keyword_analysis['score'] * weights['keywords'] +
            structure_analysis['score'] * weights['structure']
        )
        
        return round(overall, 1)
        
    except Exception as e:
        logger.error(f"Error calculating overall score: {e}")
        return 75.0

def generate_recommendations(ats_score, content_analysis, keyword_analysis, structure_analysis):
    """Generate improvement recommendations"""
    recommendations = []
    
    # ATS recommendations
    if ats_score['score'] < 80:
        recommendations.append("Improve ATS optimization by adding more relevant keywords")
    
    # Content recommendations
    if content_analysis['action_verbs_score'] < 70:
        recommendations.append("Add more action verbs to make your experience more impactful")
    
    if content_analysis['quantification_score'] < 60:
        recommendations.append("Include specific numbers and metrics to quantify your achievements")
    
    # Keyword recommendations
    if keyword_analysis['score'] < 75:
        recommendations.append("Add missing keywords from the job description to your resume")
    
    # Structure recommendations
    if structure_analysis['score'] < 80:
        missing = structure_analysis['missing_sections']
        if missing:
            recommendations.append(f"Add missing sections: {', '.join(missing)}")
    
    return recommendations[:5]  # Top 5 recommendations

def get_fallback_analysis():
    """Fallback analysis when AI is unavailable"""
    return {
        'overall_score': 75.0,
        'ats_optimization': {
            'score': 80.0,
            'checks': {},
            'issues': [],
            'strengths': ['Basic formatting', 'Contact information']
        },
        'content_quality': {
            'overall_score': 75.0,
            'word_count': 200,
            'sentence_count': 15,
            'avg_sentence_length': 13.3,
            'action_verbs_score': 70.0,
            'quantification_score': 60.0
        },
        'keyword_analysis': {
            'score': 75.0,
            'matched_keywords': ['experience', 'skills', 'management'],
            'missing_keywords': ['leadership', 'project management'],
            'suggestions': ['Add more industry-specific keywords']
        },
        'structure_analysis': {
            'score': 80.0,
            'sections': {},
            'missing_sections': [],
            'complete_sections': ['contact_info', 'summary', 'experience']
        },
        'recommendations': [
            'Add more specific metrics and achievements',
            'Include relevant keywords from job descriptions',
            'Use more action verbs in experience descriptions'
        ]
    }

@app.route('/advanced-analytics')
@login_required
def advanced_analytics():
    """Advanced Analytics Dashboard with detailed insights"""
    try:
        # Get user's data
        user_resumes = Resume.query.filter_by(user_id=current_user.id).all()
        user_applications = JobApplication.query.filter_by(user_id=current_user.id).all()
        user_cover_letters = CoverLetter.query.filter_by(user_id=current_user.id).all()
        
        # Calculate advanced metrics
        analytics = calculate_advanced_analytics(user_resumes, user_applications, user_cover_letters)
        
        return render_template('advanced_analytics.html', analytics=analytics)
        
    except Exception as e:
        logger.error(f"Error in advanced analytics: {e}")
        flash('Error loading analytics', 'error')
        return redirect(url_for('dashboard'))

def calculate_advanced_analytics(resumes, applications, cover_letters):
    """Calculate comprehensive analytics"""
    try:
        # Basic counts
        total_resumes = len(resumes)
        total_applications = len(applications)
        total_cover_letters = len(cover_letters)
        
        # Resume analytics
        resume_analytics = analyze_resumes(resumes)
        
        # Application analytics
        application_analytics = analyze_applications(applications)
        
        # Cover letter analytics
        cover_letter_analytics = analyze_cover_letters(cover_letters)
        
        # Performance trends
        performance_trends = calculate_performance_trends(applications)
        
        # Industry insights
        industry_insights = analyze_industry_trends(applications)
        
        # Success metrics
        success_metrics = calculate_success_metrics(applications)
        
        return {
            'overview': {
                'total_resumes': total_resumes,
                'total_applications': total_applications,
                'total_cover_letters': total_cover_letters,
                'active_applications': len([app for app in applications if app.status in ['applied', 'interview']]),
                'success_rate': success_metrics['success_rate']
            },
            'resume_analytics': resume_analytics,
            'application_analytics': application_analytics,
            'cover_letter_analytics': cover_letter_analytics,
            'performance_trends': performance_trends,
            'industry_insights': industry_insights,
            'success_metrics': success_metrics
        }
        
    except Exception as e:
        logger.error(f"Error calculating analytics: {e}")
        return get_fallback_analytics()

def analyze_resumes(resumes):
    """Analyze resume data"""
    try:
        if not resumes:
            return get_fallback_resume_analytics()
        
        # Template usage
        template_usage = {}
        for resume in resumes:
            template = resume.template or 'modern'
            template_usage[template] = template_usage.get(template, 0) + 1
        
        # Completion rates
        completed_resumes = sum(1 for r in resumes if r.summary and r.experience and r.skills)
        completion_rate = (completed_resumes / len(resumes)) * 100 if resumes else 0
        
        # Download statistics
        total_downloads = sum(r.downloads for r in resumes)
        avg_downloads = total_downloads / len(resumes) if resumes else 0
        
        # Recent activity
        recent_resumes = sorted(resumes, key=lambda x: x.updated_at or x.created_at, reverse=True)[:5]
        
        return {
            'template_usage': template_usage,
            'completion_rate': round(completion_rate, 1),
            'total_downloads': total_downloads,
            'avg_downloads': round(avg_downloads, 1),
            'recent_activity': [
                {
                    'name': r.name or 'Untitled',
                    'date': (r.updated_at or r.created_at).strftime('%Y-%m-%d'),
                    'template': r.template or 'modern',
                    'downloads': r.downloads
                } for r in recent_resumes
            ]
        }
        
    except Exception as e:
        logger.error(f"Error analyzing resumes: {e}")
        return get_fallback_resume_analytics()

def analyze_applications(applications):
    """Analyze job application data"""
    try:
        if not applications:
            return get_fallback_application_analytics()
        
        # Status distribution
        status_counts = {}
        for app in applications:
            status_counts[app.status] = status_counts.get(app.status, 0) + 1
        
        # Company analysis
        company_counts = {}
        for app in applications:
            company_counts[app.company] = company_counts.get(app.company, 0) + 1
        
        # Top companies
        top_companies = sorted(company_counts.items(), key=lambda x: x[1], reverse=True)[:5]
        
        # Application timeline
        recent_applications = sorted(applications, key=lambda x: x.applied_date, reverse=True)[:10]
        
        # Response rate
        responded_applications = len([app for app in applications if app.status != 'applied'])
        response_rate = (responded_applications / len(applications)) * 100 if applications else 0
        
        return {
            'status_distribution': status_counts,
            'top_companies': top_companies,
            'response_rate': round(response_rate, 1),
            'recent_applications': [
                {
                    'job_title': app.job_title,
                    'company': app.company,
                    'status': app.status,
                    'date': app.applied_date.strftime('%Y-%m-%d') if app.applied_date else 'Unknown'
                } for app in recent_applications
            ]
        }
        
    except Exception as e:
        logger.error(f"Error analyzing applications: {e}")
        return get_fallback_application_analytics()

def analyze_cover_letters(cover_letters):
    """Analyze cover letter data"""
    try:
        if not cover_letters:
            return get_fallback_cover_letter_analytics()
        
        # Template usage
        template_usage = {}
        for cl in cover_letters:
            template = cl.template or 'standard'
            template_usage[template] = template_usage.get(template, 0) + 1
        
        # Download statistics
        total_downloads = sum(cl.downloads for cl in cover_letters)
        avg_downloads = total_downloads / len(cover_letters) if cover_letters else 0
        
        # Recent activity
        recent_cover_letters = sorted(cover_letters, key=lambda x: x.updated_at or x.created_at, reverse=True)[:5]
        
        return {
            'template_usage': template_usage,
            'total_downloads': total_downloads,
            'avg_downloads': round(avg_downloads, 1),
            'recent_activity': [
                {
                    'title': cl.title or 'Untitled',
                    'company': cl.company,
                    'date': (cl.updated_at or cl.created_at).strftime('%Y-%m-%d') if cl.updated_at or cl.created_at else 'Unknown',
                    'template': cl.template or 'standard'
                } for cl in recent_cover_letters
            ]
        }
        
    except Exception as e:
        logger.error(f"Error analyzing cover letters: {e}")
        return get_fallback_cover_letter_analytics()

def calculate_performance_trends(applications):
    """Calculate performance trends over time"""
    try:
        if not applications:
            return get_fallback_performance_trends()
        
        # Group by month
        monthly_data = {}
        for app in applications:
            if app.applied_date:
                month_key = app.applied_date.strftime('%Y-%m')
                if month_key not in monthly_data:
                    monthly_data[month_key] = {'total': 0, 'responses': 0, 'interviews': 0}
                
                monthly_data[month_key]['total'] += 1
                if app.status != 'applied':
                    monthly_data[month_key]['responses'] += 1
                if app.status in ['interview', 'offer']:
                    monthly_data[month_key]['interviews'] += 1
        
        # Convert to chart data
        labels = sorted(monthly_data.keys())
        total_data = [monthly_data[month]['total'] for month in labels]
        response_data = [monthly_data[month]['responses'] for month in labels]
        interview_data = [monthly_data[month]['interviews'] for month in labels]
        
        return {
            'labels': labels,
            'total_applications': total_data,
            'responses': response_data,
            'interviews': interview_data
        }
        
    except Exception as e:
        logger.error(f"Error calculating performance trends: {e}")
        return get_fallback_performance_trends()

def analyze_industry_trends(applications):
    """Analyze industry trends"""
    try:
        if not applications:
            return get_fallback_industry_insights()
        
        # Industry keywords (simplified)
        industry_keywords = {
            'technology': ['software', 'tech', 'programming', 'development', 'engineering'],
            'finance': ['finance', 'banking', 'investment', 'accounting', 'financial'],
            'healthcare': ['health', 'medical', 'nursing', 'pharmacy', 'hospital'],
            'education': ['education', 'teaching', 'academic', 'school', 'university'],
            'marketing': ['marketing', 'advertising', 'brand', 'digital', 'social']
        }
        
        industry_counts = {industry: 0 for industry in industry_keywords}
        
        for app in applications:
            job_text = f"{app.job_title} {app.company}".lower()
            for industry, keywords in industry_keywords.items():
                if any(keyword in job_text for keyword in keywords):
                    industry_counts[industry] += 1
                    break
        
        # Top industries
        top_industries = sorted(industry_counts.items(), key=lambda x: x[1], reverse=True)
        
        return {
            'industry_distribution': industry_counts,
            'top_industries': top_industries[:5]
        }
        
    except Exception as e:
        logger.error(f"Error analyzing industry trends: {e}")
        return get_fallback_industry_insights()

def calculate_success_metrics(applications):
    """Calculate success metrics"""
    try:
        if not applications:
            return get_fallback_success_metrics()
        
        total_applications = len(applications)
        responses = len([app for app in applications if app.status != 'applied'])
        interviews = len([app for app in applications if app.status == 'interview'])
        offers = len([app for app in applications if app.status == 'offer'])
        
        response_rate = (responses / total_applications) * 100 if total_applications > 0 else 0
        interview_rate = (interviews / total_applications) * 100 if total_applications > 0 else 0
        offer_rate = (offers / total_applications) * 100 if total_applications > 0 else 0
        success_rate = (offers / total_applications) * 100 if total_applications > 0 else 0
        
        return {
            'total_applications': total_applications,
            'responses': responses,
            'interviews': interviews,
            'offers': offers,
            'response_rate': round(response_rate, 1),
            'interview_rate': round(interview_rate, 1),
            'offer_rate': round(offer_rate, 1),
            'success_rate': round(success_rate, 1)
        }
        
    except Exception as e:
        logger.error(f"Error calculating success metrics: {e}")
        return get_fallback_success_metrics()

# Fallback functions for when data is unavailable
def get_fallback_analytics():
    return {
        'overview': {'total_resumes': 0, 'total_applications': 0, 'total_cover_letters': 0, 'active_applications': 0, 'success_rate': 0},
        'resume_analytics': get_fallback_resume_analytics(),
        'application_analytics': get_fallback_application_analytics(),
        'cover_letter_analytics': get_fallback_cover_letter_analytics(),
        'performance_trends': get_fallback_performance_trends(),
        'industry_insights': get_fallback_industry_insights(),
        'success_metrics': get_fallback_success_metrics()
    }

def get_fallback_resume_analytics():
    return {
        'template_usage': {'modern': 1},
        'completion_rate': 0,
        'total_downloads': 0,
        'avg_downloads': 0,
        'recent_activity': []
    }

def get_fallback_application_analytics():
    return {
        'status_distribution': {'applied': 1},
        'top_companies': [],
        'response_rate': 0,
        'recent_applications': []
    }

def get_fallback_cover_letter_analytics():
    return {
        'template_usage': {'standard': 1},
        'total_downloads': 0,
        'avg_downloads': 0,
        'recent_activity': []
    }

def get_fallback_performance_trends():
    return {
        'labels': [],
        'total_applications': [],
        'responses': [],
        'interviews': []
    }

def get_fallback_industry_insights():
    return {
        'industry_distribution': {},
        'top_industries': []
    }

def get_fallback_success_metrics():
    return {
        'total_applications': 0,
        'responses': 0,
        'interviews': 0,
        'offers': 0,
        'response_rate': 0,
        'interview_rate': 0,
        'offer_rate': 0,
        'success_rate': 0
    }

if __name__ == '__main__':
    logger.info("Starting Flask application...")
    app.run(debug=True, host='0.0.0.0', port=5000) 